# dashboard_mapped.py
import streamlit as st
import pandas as pd
import plotly.express as px
import re
import geopandas as gpd
import pydeck as pdk
import os
import glob
from PIL import Image
from io import BytesIO
import base64
from streamlit_plotly_events import plotly_events
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import requests
from streamlit import cache_data
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from collections import OrderedDict
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.styles import Border, Side
import io
from io import BytesIO
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import numbers

# --- Page config for wide layout ---
st.set_page_config(
    page_title="Gaeltec Dashboard",
    layout="wide",  # <-- makes the dashboard wider
    initial_sidebar_state="expanded"
)

def sanitize_sheet_name(name: str) -> str:
    """
    Remove or replace invalid characters for Excel sheet names.
    Excel sheet names cannot contain: : \ / ? * [ ]
    """
    name = str(name)
    name = re.sub(r'[:\\/*?\[\]]', '_', name)
    name = re.sub(r'[^\x00-\x7F]', '_', name)
    return name[:31]

def get_scottish_weather(api_key, location="Ayrshire"):
    """
    Get weather data for Scottish locations
    """
    # Coordinates for Scottish locations
    locations = {
        "Ayrshire": {"lat": 55.458, "lon": -4.629},
        "Lanarkshire": {"lat": 55.676, "lon": -3.785},
        "Glasgow": {"lat": 55.864, "lon": -4.252},
        "Edinburgh": {"lat": 55.953, "lon": -3.188}
    }
    
    if location in locations:
        coords = locations[location]
    else:
        # Default to Ayrshire
        coords = locations["Ayrshire"]
    
    base_url = "http://api.openweathermap.org/data/2.5/weather"
    params = {
        'lat': coords["lat"],
        'lon': coords["lon"],
        'appid': api_key,
        'units': 'metric'
    }
    
    try:
        response = requests.get(base_url, params=params)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        st.error(f"Error fetching weather data: {e}")
        return None

@cache_data(ttl=1800)  # Cache for 30 minutes
def get_weather_forecast(api_key, location="Ayrshire"):
    """
    Get 5-day forecast for Scottish locations
    """
    locations = {
        "Ayrshire": {"lat": 55.458, "lon": -4.629},
        "Lanarkshire": {"lat": 55.676, "lon": -3.785}
    }
    
    if location in locations:
        coords = locations[location]
    else:
        coords = locations["Ayrshire"]
    
    base_url = "http://api.openweathermap.org/data/2.5/forecast"
    params = {
        'lat': coords["lat"],
        'lon': coords["lon"],
        'appid': api_key,
        'units': 'metric'
    }
    
    try:
        response = requests.get(base_url, params=params)
        response.raise_for_status()
        return response.json()
    except Exception as e:
        st.error(f"Forecast API error: {e}")
        return None


def poles_to_word(df: pd.DataFrame) -> BytesIO:
    doc = Document()

    # Defensive cleaning
    df = df.copy()
    df = df.replace(
        to_replace=["nan", "NaN", "None", None],
        value=""
    )

    grouped = df.groupby('pole', sort=False)

    for pole, group in grouped:
        pole_str = str(pole).strip()
        if not pole_str:
            continue

        # Ordered set using dict keys (preserves order, removes duplicates)
        unique_texts = OrderedDict()

        for _, row in group.iterrows():
            parts = []

            wi = str(row.get('Work instructions', '')).strip()
            comment = str(row.get('comment', '')).strip()

            if wi:
                parts.append(wi)

            if comment:
                parts.append(f"({comment})")

            if parts:
                text = " ".join(parts)

                # Normalize for deduplication
                normalized = text.lower().strip()

                unique_texts[normalized] = text

        if not unique_texts:
            continue

        # Bullet paragraph
        p = doc.add_paragraph(style='List Bullet')

        run_number = p.add_run(f"{pole_str} – ")
        run_number.bold = True
        run_number.font.name = 'Times New Roman'
        run_number.font.size = Pt(12)

        texts = list(unique_texts.values())

        for i, text in enumerate(texts):
            run_item = p.add_run(text)
            run_item.bold = True
            run_item.font.name = 'Times New Roman'
            run_item.font.size = Pt(12)

            if "Erect Pole" in text:
                run_item.font.highlight_color = WD_COLOR_INDEX.RED

            if i < len(texts) - 1:
                p.add_run(" ; ")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def build_export_df(filtered_df):
    export_df = filtered_df.copy()

    # Rename columns
    export_df = export_df.rename(columns=column_rename_map)

    # Keep only columns that actually exist
    existing_cols = [c for c in export_columns if c in export_df.columns]
    export_df = export_df[existing_cols]

    return export_df

# Normalize strings: remove leading/trailing spaces, lowercase, remove extra dots
def normalize_item(s):
    if pd.isna(s):
        return ""
    s = str(s).strip().lower()           # strip spaces and lowercase
    s = s.replace(".", "")               # remove dots
    s = re.sub(r"\s+", " ", s)          # collapse multiple spaces
    return s

def apply_common_filters(df):
    df = df.copy()

    # Ensure datetime
    df['datetouse_dt'] = pd.to_datetime(df['datetouse'], errors='coerce')

    # Date rule: after 2023
    df = df[df['datetouse_dt'].dt.year > 2023]

    # Segment
    if selected_segment != 'All' and 'segmentcode' in df.columns:
        df = df[
            df['segmentcode'].astype(str).str.strip()
            == str(selected_segment).strip()
        ]

    # Pole
    if selected_pole != "All" and 'pole' in df.columns:
        df = df[
            df['pole'].astype(str).str.strip()
            == str(selected_pole).strip()
        ]

    # Ensure numeric total
    if 'total' in df.columns:
        df['total'] = pd.to_numeric(df['total'], errors='coerce')

    return df.dropna(subset=['datetouse_dt'])
    
def prepare_dataframe(df):
    df = df.copy()
    df.columns = df.columns.str.strip().str.lower()

    if 'datetouse' in df.columns:
        df['datetouse_dt'] = pd.to_datetime(df['datetouse'], errors='coerce').dt.normalize()
    else:
        df['datetouse_dt'] = pd.NaT

    # Make numeric columns safe
    for col in ['total', 'orig']:
        if col in df.columns:
            df[col] = (
                df[col].astype(str)
                .str.replace(" ", "")
                .str.replace(",", ".", regex=False)
                .astype(float)
            )

    return df

def multi_select_filter(col, label, df):
    if col not in df.columns:
        return ["All"], df

    options = ["All"] + sorted(df[col].dropna().astype(str).unique())
    selected = st.sidebar.multiselect(label, options, default=["All"])

    if "All" in selected:
        return selected, df

    return selected, df[df[col].astype(str).isin(selected)]



def to_excel(project_df, team_df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:

        # ---- Sheet 1: Revenue per Project ----
        if not project_df.empty:
            project_df.to_excel(writer, index=False, sheet_name="Revenue per Project", startrow=1)
            ws_proj = writer.sheets["Revenue per Project"]

            # ---- Column widths ----
            ws_proj.column_dimensions["A"].width = 30
            ws_proj.column_dimensions["B"].width = 18

            # ---- Styles ----
            header_font = Font(bold=True, size=14)
            header_fill = PatternFill(start_color="00CCFF", end_color="00CCFF", fill_type="solid")
            thin_side = Side(style="thin")
            medium_side = Side(style="medium")
            thick_side = Side(style="thick")
            light_grey_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
            white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

            max_col = ws_proj.max_column
            max_row = ws_proj.max_row

            # Set row 1 height for images
            ws_proj.row_dimensions[1].height = 120

            # Header → row 2
            for col_idx, cell in enumerate(ws_proj[2], start=1):
                cell.font = header_font
                cell.fill = header_fill
                cell.border = Border(
                    left=thick_side if col_idx == 1 else medium_side,
                    right=thick_side if col_idx == max_col else medium_side,
                    top=thick_side,
                    bottom=thick_side
                )

            # Data rows → start row 3
            for row_idx in range(3, max_row + 1):
                fill = light_grey_fill if row_idx % 2 == 0 else white_fill
                for col_idx in range(1, max_col + 1):
                    cell = ws_proj.cell(row=row_idx, column=col_idx)
                    cell.fill = fill
                    cell.border = Border(
                        left=thin_side,
                        right=thin_side,
                        top=thin_side,
                        bottom=thin_side
                    )

            # ---- Add images in row 1 ----
            img1 = XLImage("Images/GaeltecImage.png")
            img2 = XLImage("Images/SPEN.png")
            img1.width = 120; img1.height = 120; img1.anchor = "A1"
            img2.width = 360; img2.height = 120; img2.anchor = "B1"
            ws_proj.add_image(img1)
            ws_proj.add_image(img2)

        # ---- Sheet 2: Revenue per Team ----
        if not team_df.empty:
            team_df.to_excel(writer, index=False, sheet_name="Revenue per Team", startrow=1)
            ws_team = writer.sheets["Revenue per Team"]

            ws_team.column_dimensions["A"].width = 25
            ws_team.column_dimensions["B"].width = 18

            max_col = ws_team.max_column
            max_row = ws_team.max_row

            # Row 1 for images
            ws_team.row_dimensions[1].height = 120

            # Header → row 2
            for col_idx, cell in enumerate(ws_team[2], start=1):
                cell.font = header_font
                cell.fill = header_fill
                cell.border = Border(
                    left=thick_side if col_idx == 1 else medium_side,
                    right=thick_side if col_idx == max_col else medium_side,
                    top=thick_side,
                    bottom=thick_side
                )

            # Data rows → start row 3
            for row_idx in range(3, max_row + 1):
                fill = light_grey_fill if row_idx % 2 == 0 else white_fill
                for col_idx in range(1, max_col + 1):
                    cell = ws_team.cell(row=row_idx, column=col_idx)
                    cell.fill = fill
                    cell.border = Border(
                        left=thin_side,
                        right=thin_side,
                        top=thin_side,
                        bottom=thin_side
                    )

            # ---- Add images in row 1 ----
            img1 = XLImage("Images/GaeltecImage.png")
            img2 = XLImage("Images/SPEN.png")
            img1.width = 120; img1.height = 120; img1.anchor = "A1"
            img2.width = 360; img2.height = 120; img2.anchor = "B1"
            ws_team.add_image(img1)
            ws_team.add_image(img2)

    output.seek(0)
    return output

def generate_excel_styled_multilevel(filtered_df, poles_df=None):
    wb = Workbook()
    ws = wb.active
    ws.title = "Daily Revenue"

    # ---- Sheet 1: Daily Revenue ----
    if {'shire', 'project', 'segmentcode', 'projectmanager', 'datetouse_dt', 'total'}.issubset(filtered_df.columns):
        daily_df = (
            filtered_df
            .groupby(['datetouse_dt','shire','project','segmentcode','projectmanager'], as_index=False)
            .agg({'total':'sum'})
        )
        daily_df.rename(columns={
            'datetouse_dt':'Date',
            'total':'Revenue (£)',
            'segmentcode':'Segment',
            'projectmanager':'Project Manager'
        }, inplace=True)

        # Write header in ROW 2 (row 1 reserved for images)
        for col_idx, col_name in enumerate(daily_df.columns.tolist(), start=1):
            ws.cell(row=2, column=col_idx, value=col_name)

        # Write data starting from row 3
        for r_idx, row in enumerate(daily_df.values.tolist(), start=3):
            for c_idx, value in enumerate(row, start=1):
                ws.cell(row=r_idx, column=c_idx, value=value)

    # ---- Sheet 2: Poles Summary ----
    ws_summary = wb.create_sheet(title="Poles Summary")
    if poles_df is not None and not poles_df.empty:
        poles_summary = (
            poles_df[['shire','project','segmentcode','pole']]
            .drop_duplicates()
            .groupby(['shire','project','segmentcode'], as_index=False)
            .agg({'pole': lambda x: ', '.join(sorted(x.astype(str)))})
        )
        poles_summary.rename(columns={'pole':'Poles', 'segmentcode':'Segment'}, inplace=True)

        # Write multi-level headers (Row 2-4)
        headers = ['Shire','Project','Segment','Poles']
        for idx, h in enumerate(headers, start=1):
            ws_summary.cell(row=2, column=idx, value=h)  # Shire header
            ws_summary.cell(row=3, column=idx, value=h if h != 'Poles' else '')  # Project header
            ws_summary.cell(row=4, column=idx, value=h if h != 'Poles' else '')  # Segment header

        # Write data starting from row 5
        for r_idx, row in enumerate(poles_summary.values.tolist(), start=5):
            for c_idx, value in enumerate(row, start=1):
                ws_summary.cell(row=r_idx, column=c_idx, value=value)

    # ---- Formatting styles ----
    header_font = Font(bold=True, size=16)
    header_fill = PatternFill(start_color="00CCFF", end_color="00CCFF", fill_type="solid")
    thin_side = Side(style="thin")
    medium_side = Side(style="medium")
    thick_side = Side(style="thick")
    light_grey_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

    # ---- Add images ----
    IMG_HEIGHT = 120
    IMG_WIDTH_SMALL = 120
    IMG_WIDTH_LARGE = IMG_WIDTH_SMALL * 3

    # Set row 1 height to fit images
    ws.row_dimensions[1].height = IMG_HEIGHT * 0.75  # approximate pixels → Excel points
    ws_summary.row_dimensions[1].height = IMG_HEIGHT * 0.75

    # Position images (row 1)
    img1.anchor = "B1"
    img2.anchor = "A1"

    ws.add_image(img1)
    ws.add_image(img2)

    # Same for Summary
    img1_s = XLImage("Images/GaeltecImage.png")
    img2_s = XLImage("Images/SPEN.png")

    img1_s.width = IMG_WIDTH_SMALL
    img1_s.height = IMG_HEIGHT
    img1_s.anchor = "A1"

    img2_s.width = IMG_WIDTH_LARGE
    img2_s.height = IMG_HEIGHT
    img2_s.anchor = "B1"

    # Sheet 2 images
    img1_s = XLImage("Images/GaeltecImage.png")
    img2_s = XLImage("Images/SPEN.png")
    img1_s.width = IMG_WIDTH_SMALL; img1_s.height = IMG_HEIGHT; img1_s.anchor = "A1"
    img2_s.width = IMG_WIDTH_LARGE; img2_s.height = IMG_HEIGHT; img2_s.anchor = "B1"
    ws_summary.add_image(img1_s)
    ws_summary.add_image(img2_s)

    # ---- Apply formatting ----
    for sheet in [ws, ws_summary]:
        max_col = sheet.max_column
        max_row = sheet.max_row

        # Header rows
        for row_idx in range(2, 5 if sheet == ws_summary else 3):
            for col_idx in range(1, max_col + 1):
                cell = sheet.cell(row=row_idx, column=col_idx)
                cell.font = header_font
                cell.fill = header_fill
                sheet.column_dimensions[get_column_letter(col_idx)].width = 60 if col_idx == 1 else 20
                cell.border = Border(
                    left=thick_side if col_idx == 1 else medium_side,
                    right=thick_side if col_idx == max_col else medium_side,
                    top=thick_side,
                    bottom=thick_side
                )

        # DATA ROWS → after headers
        start_data_row = 5 if sheet == ws_summary else 3
        for row_idx in range(start_data_row, max_row + 1):
            fill = light_grey_fill if row_idx % 2 == 1 else white_fill
            for col_idx in range(1, max_col + 1):
                cell = sheet.cell(row=row_idx, column=col_idx)
                cell.fill = fill
                cell.border = Border(
                    left=thin_side,
                    right=thin_side,
                    top=thin_side,
                    bottom=thin_side
                )

    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

    
# --- MAPPINGS ---

# --- Project Manager Mapping ---
project_mapping = {
    "Jonathon Mcclung": ["Ayrshire", "PCB"],
    "Gary MacDonald": ["Ayrshire", "LV"],
    "Jim Gaffney": ["Lanark", "PCB"],
    "Calum Thomson": ["Ayrshire", "Connections"],
    "Calum Thomsom": ["Ayrshire", "Connections"],
    "Calum Thompson": ["Ayrshire", "Connections"],
    "Andrew Galt": ["Ayrshire", "-"],
    "Henry Gordon": ["Ayrshire", "-"],
    "Jonathan Douglas": ["Ayrshire", "11 kV"],
    "Jonathon Douglas": ["Ayrshire", "11 kV"],
    "Matt": ["Lanark", ""],
    "Lee Fraser": ["Ayrshire", "Connections"],
    "Lee Frazer": ["Ayrshire", "Connections"],
    "Mark": ["Lanark", "Connections"],
    "Mark Nicholls": ["Ayrshire", "Connections"],
    "Cameron Fleming": ["Lanark", "Connections"],
    "Ronnie Goodwin": ["Lanark", "Connections"],
    "Ian Young": ["Ayrshire", "Connections"],
    "Matthew Watson": ["Lanark", "Connections"],
    "Aileen Brese": ["Ayrshire", "Connections"],
    "Mark McGoldrick": ["Lanark", "Connections"]
}

# --- Region Mapping ---
mapping_region = {
    "Newmilns": ["Irvine Valley"],
    "New Cumnock": ["New Cumnock"],
    "Kilwinning": ["Kilwinning"],
    "Stewarton": ["Irvine Valley"],
    "Kilbirnie": ["Kilbirnie and Beith"],
    "Coylton": ["Ayr East"],
    "Irvine": ["Irvine Valley", "Irvine East", "Irvine West"],
    "TROON": ["Troon"],
    "Ayr": ["Ayr East", "Ayr North", "Ayr West"],
    "Maybole": ["Maybole, North Carrick and Coylton"],
    "Clerkland": ["Irvine Valley"],
    "Glengarnock": ["Kilbirnie and Beith"],
    "Ayrshire": ["North Coast and Cumbraes","Prestwick", "Saltcoats and Stevenston", "Troon", "Ayr East", "Ayr North",
                 "Ayr West","Annick","Ardrossan and Arran","Dalry and West Kilbride","Girvan and South Carrick","Irvine East",
                 "Irvine Valley","Irvine West","Kilbirnie and Beith","Kilmarnock East and Hurlford","Kilmarnock North",
                 "Kilmarnock South","Kilmarnock West and Crosshouse","Kilwinning","Kyle","Maybole, North Carrick and Coylton",
                 "Ayr, Carrick and Cumnock","East_Ayrshire","North_Ayrshre","South_Ayrshre","Doon Valley"],
    "Lanark": ["Abronhill, Kildrum and the Village","Airdrie Central","Airdrie North","Airdrie South","Avondale and Stonehouse",
               "Ballochmyle","Bellshill","Blantyre","Bothwell and Uddingston","Cambuslang East","Cambuslang West",
               "Clydesdale East","Clydesdale North","Clydesdale South","Clydesdale West","Coatbridge North and Glenboig",
               "Coatbridge South","Coatbridge West","Cumbernauld North","Cumbernauld South",
               "East Kilbride Central North","East Kilbride Central South","East Kilbride East","East Kilbride South",
               "East Kilbride West","Fortissat","Hamilton North and East","Hamilton South","Hamilton West and Earnock",
               "Mossend and Holytown","Motherwell North","Motherwell South East and Ravenscraig","Motherwell West",
               "Rutherglen Central and North","Rutherglen South","Strathkelvin","Thorniewood","Wishaw","Larkhall",
               "Airdrie and Shotts","Cumbernauld, Kilsyth and Kirkintilloch East","East Kilbride, Strathaven and Lesmahagow",
               "Lanark and Hamilton East","Motherwell and Wishaw","North_Lanarkshire","South_Lanarkshire"]
}

# --- File Project Mapping ---
file_project_mapping = {
    "pcb 2022": ["Ayrshire", "PCB"],
    "33kv refurb": ["Ayrshire", "33kv Refurb"],
    "connections": ["Ayrshire", "Connections"],
    "storms": ["Ayrshire", "Storms"],
    "11kv refurb": ["Ayrshire", "11kv Refurb"],
    "11kV Refurb Ayrshire 2026": ["Ayrshire", "11kV Refurb"],
    "11kV Refurb Ayrshire Pinwherry": ["Ayrshire", "11kV Refurb"],
    "aurs road": ["Ayrshire", "Aurs Road"],
    "spen labour": ["Ayrshire", "SPEN Labour"],
    "lvhi5": ["Ayrshire", "LV"],
    "pcb": ["Ayrshire", "PCB"],
    "lanark": ["Lanark", ""],
    "11kv refur": ["Lanark", "11kv Refurb"],
    "lv & esqcr": ["Lanark", "LV"],
    "11kv rebuilt": ["Lanark", "11kV Rebuilt"],
    "33kv rebuilt": ["Lanark", "33kV Rebuilt"]
}

# --- Pole Mappings (dictionary style, includes new additions) ---
pole_keys = {
    "9x220 BIOCIDE LV POLE": "9m B",
    "9x275 BIOCIDE LV POLE": "9s B",
    "9x220 CREOSOTE LV POLE": "9m",
    "9x275 CREOSOTE LV POLE": "9s",
    "9x220 HV SINGLE POLE": "9m",
    "9x275 HV SINGLE POLE": "9s",
    "9x295 HV SINGLE POLE": "9es",
    "9x315 HV SINGLE POLE": "9esp",
    "10x230 BIOCIDE LV POLE": "10m B",
    "10x230 HV SINGLE POLE": "10m",
    "10x285 BIOCIDE LV POLE": "10s B",
    "10x285 H POLE HV Creosote": "10s",
    "10x285 HV SINGLE POLE": "10s",
    "10x305 HV SINGLE POLE": "10es",
    "11x295 HV SINGLE POLE": "11s",
    "11x295 H POLE HV Creosote": "11s",
    "11x295 BIOCIDE LV POLE": "11sB",
    "12x250 BIOCIDE LV POLE": "12m B",
    "12x305 BIOCIDE LV POLE": "12s B",
    "12x250 CREOSOTE LV POLE": "12m",
    "12x305 CREOSOTE LV POLE": "12s",
    "12x305 H POLE HV Creosote":"12s",
    "12x250 HV SINGLE POLE": "12m",
    "12x305 HV SINGLE POLE": "12s",
    "12x325 HV SINGLE POLE": "12es",
    "12x345 HV SINGLE POLE": "12esp",
    "13x260 BIOCIDE LV POLE": "13m B",
    "13x320 BIOCIDE LV POLE": "13s B",
    "13x260 CREOSOTE LV POLE": "13m",
    "13x320 CREOSOTE LV POLE": "13s",
    "13x260 HV SINGLE POLE": "13m",
    "13x320 HV SINGLE POLE": "13s",
    "13x340 HV SINGLE POLE": "13es",
    "13x365 HV SINGLE POLE": "13esp",
    "14x275 BIOCIDE LV POLE": "14m B",
    "14x335 BIOCIDE LV POLE": "14s B",
    "14x275 CREOSOTE LV POLE": "14m",
    "14x335 CREOSOTE LV POLE": "14s",
    "14x275 HV SINGLE POLE": "14m",
    "14x335 HV SINGLE POLE": "14s",
    "14x355 HV SINGLE POLE": "14es",
    "14x375 HV SINGLE POLE": "14esp",
    "16x305 BIOCIDE LV POLE": "16m B",
    "16x365 BIOCIDE LV POLE": "16s B",
    "16x305 CREOSOTE LV POLE": "16m",
    "16x365 CREOSOTE LV POLE": "16s",
    "16x305 HV SINGLE POLE": "16m",
    "16x365 HV SINGLE POLE": "16s",
    "16x385 HV SINGLE POLE": "16es",
    "16x405 HV SINGLE POLE": "16esp",
    "11x315 H POLE HV Creosote":"11es",
    "14x335 H POLE HV Creosote":"14s",
    "11x315 HV SINGLE POLE":"11es",
    "13x320 H POLE HV Creosote":"13s",
    "11x240 CREOSOTE LV POLE":"11",
    "11x240 HV SINGLE POLE":"11m",
    "10x230 CREOSOTE LV POLE":"10m",
    "11x335 H POLE HV Creosote":"11esp",
    "10x305 H POLE HV Creosote":"10es",
    "11x240 BIOCIDE LV POLE":"11m B",
    "16x365 H POLE HV Creosote":"16s",
    "16x405 EHV SINGLE POLE CREOSOTE":"16esp",
    "12x325 H POLE HV Creosote":"12es",
    "16x385 H POLE HV Creosote":"16es",
    "12x305 EHV SINGLE POLE CREOSOTE":"12s",
    "13x340 EHV SINGLE POLE CREOSOTE":"13es",
    "11x335 EHV SINGLE POLE CREOSOTE":"11es",
    "11x315 EHV SINGLE POLE CREOSOTE":"11es",
    "12x325 EHV SINGLE POLE CREOSOTE":"12es"
}

pole_erected_keys = {
    "Erect Single HV/EHV Pole, up to and including 12 metre pole":"Erect HV pole", 
    "Erect Single HV/EHV Pole, up to and including 12 metre pole.":"Erect HV pole",
    "Erect LV Structure Single Pole, up to and including 12 metre pole" :"Erect LV pole",
    "Erect Section Structure 'H' HV/EHV Pole, up to and including 12 metre pole.":"H HV pole"
}

poles_replaced_keys = {
    "Recover single pole, up to and including 15 metres in height, and reinstate, all ground conditions":"Recover single pole",
    "Recover 'A' / 'H' pole, up to and including 15 metres in height, and reinstate, all ground conditions":"Recover H pole"
}

# --- Equipment / Conductor Mappings ---
equipment_keys = {
    "Hazel - 50mm² AAAC bare (1000m drums)": "Hazel 50mm² (1000m drums)",
    "Oak - 100mm² AAAC bare (1000m drums)": "Oak 100mm² (1000m drums)",
    "Ash - 150mm² AAAC bare (1000m drums)": "Ash 150mm² (1000m drums)",
    "Poplar - 200mm² AAAC bare (1000m drums)": "Poplar 200mm² (1000m drums)",
    "Upas - 300mm² AAAC bare (1000m drums)": "Upas 300mm² (1000m drums)",
    "Poplar OPPC - 200mm² AAAC equivalent bare": "Poplar OPPC 200mm²",
    "Upas OPPC - 300mm² AAAC equivalent bare": "Upas OPPC 300mm²",
    # ACSR
    "Gopher - 25mm² ACSR bare (1000m drums)": "Gopher 25mm² (1000m drums)",
    "Caton - 25mm² Compacted ACSR bare (1000m drums)": "Caton 25mm² (1000m drums)",
    "Rabbit - 50mm² ACSR bare (1000m drums)": "Rabbit 50mm² (1000m drums)",
    "Wolf - 150mm² ACSR bare (1000m drums)": "Wolf 150mm² (1000m drums)",
    "Horse - 70mm² ACSR bare": "Horse 70mm²",
    "Dog - 100mm² ACSR bare (1000m drums)": "Dog 100mm² (1000m drums)",
    "Dingo - 150mm² ACSR bare (1000m drums)": "Dingo 150mm² (1000m drums)",
    # Copper
    "Hard Drawn Copper 16mm² ( 3/2.65mm ) (500m drums)": "Copper 16mm² (500m drums)",
    "Hard Drawn Copper 32mm² ( 3/3.75mm ) (1000m drums)": "Copper 32mm² (500m drums)",
    "Hard Drawn Copper 70mm² (500m drums)": "Copper 70mm² (500m drums)",
    "Hard Drawn Copper 100mm² (500m drums)": "Copper 100mm² (500m drums)",
    # PVC covered
    "35mm² Copper (Green / Yellow PVC covered) (50m drums)": "Copper 35mm² GY PVC (50m drums)",
    "70mm² Copper (Green / Yellow PVC covered) (50m drums)": "Copper 70mm² GY PVC (50m drums)",
    "35mm² Copper (Blue PVC covered) (50m drums)": "Copper 35mm² Blue PVC (50m drums)",
    "70mm² Copper (Blue PVC covered) (50m drums)": "Copper 70mm² Blue PVC (50m drums)",
    # Double insulated
    "35mm² Double Insulated (Brown) (50m drums)": "Double Insulated 35mm² Brown (50m drums)",
    "35mm² Double Insulated (Blue) (50m drums)": "Double Insulated 35mm² Blue (50m drums)",
    "70mm² Double Insulated (Brown) (50m drums)": "Double Insulated 70mm² Brown (50m drums)",
    "70mm² Double Insulated (Blue) (50m drums)": "Double Insulated 70mm² Blue (50m drums)",
    "120mm² Double Insulated (Brown) (50m drums)": "Double Insulated 120mm² Brown (50m drums)",
    "120mm² Double Insulated (Blue) (50m drums)": "Double Insulated 120mm² Blue (50m drums)",
    # LV cables
    "LV Cable 1ph 4mm Concentric (250m drums)": "LV 1ph 4mm Concentric (250m drums)",
    "LV Cable 1ph 25mm CNE (250m drums)": "LV 1ph 25mm CNE (250m drums)",
    "LV Cable 1ph 25mm SNE (100m drums)": "LV 1ph 25mm SNE (100m drums)",
    "LV Cable 1ph 35mm CNE (250m drums)": "LV 1ph 35mm CNE (250m drums)",
    "LV Cable 1ph 35mm SNE (100m drums)": "LV 1ph 35mm SNE (100m drums)",
    "LV Cable 3ph 35mm Cu Split Con (250m drums)": "LV 3ph 35mm Cu Split Con (250m drums)",
    "LV Cable 3ph 35mm SNE (250m drums)": "LV 3ph 35mm SNE (250m drums)",
    "LV Cable 3ph 35mm CNE (250m drums)": "LV 3ph 35mm CNE (250m drums)",
    "LV Cable 3ph 35mm CNE Al (LSOH) (250m drums)": "LV 3ph 35mm CNE Al LSOH (250m drums)",
    "LV Cable 3c 95mm W/F (250m drums)": "LV 3c 95mm W/F (250m drums)",
    "LV Cable 3c 185mm W/F (250m drums)": "LV 3c 185mm W/F (250m drums)",
    "LV Cable 3c 300mm W/F (250m drums)": "LV 3c 300mm W/F (250m drums)",
    "LV Cable 4c 95mm W/F (250m drums)": "LV 4c 95mm W/F (250m drums)",
    "LV Cable 4c 185mm W/F (250m drums)": "LV 4c 185mm W/F (250m drums)",
    "LV Cable 4c 240mm W/F (250m drums)": "LV 4c 240mm W/F (250m drums)",
    "LV Marker Tape (365m roll)": "LV Marker Tape (365m roll)",
    # 11kV
    "11kv Cable 95mm 3c Poly (250m drums)": "11kV 3c 95mm Poly (250m drums)",
    "11kv Cable 185mm 3c Poly (250m drums)": "11kV 3c 185mm Poly (250m drums)",
    "11kv Cable 300mm 3c Poly (250m drums)": "11kV 3c 300mm Poly (250m drums)",
    "11kv Cable 95mm 1c Poly (250m drums)": "11kV 1c 95mm Poly (250m drums)",
    "11kv Cable 185mm 1c Poly (250m drums)": "11kV 1c 185mm Poly (250m drums)",
    "11kv Cable 300mm 1c Poly (250m drums)": "11kV 1c 300mm Poly (250m drums)",
    "11kV Marker Tape (40m roll)": "11kV Marker Tape (40m roll)"
}

# --- Transformer Mappings ---
transformer_keys = {
    "Transformer 1ph 50kVA": "TX 1ph (50kVA)",
    "Transformer 3ph 50kVA": "TX 3ph (50kVA)",
    "Transformer 1ph 100kVA": "TX 1ph (100kVA)",
    "Transformer 1ph 25kVA": "TX 1ph (25kVA)",
    "Transformer 3ph 200kVA": "TX 3ph (200kVA)",
    "Transformer 3ph 100kVA": "TX 3ph (100kVA)"
}

# --- Equipment / Conductor Mappings ---
conductor_keys = {
    "Hazel - 50mm² AAAC bare (1000m drums)": "Hazel 50mm²",
    "Oak - 100mm² AAAC bare (1000m drums)": "Oak 100mm²",
    "Ash - 150mm² AAAC bare (1000m drums)": "Ash 150mm²",
    "Poplar - 200mm² AAAC bare (1000m drums)": "Poplar 200mm²",
    "Upas - 300mm² AAAC bare (1000m drums)": "Upas 300mm²",
    "Poplar OPPC - 200mm² AAAC equivalent bare": "Poplar OPPC 200mm²",
    "Upas OPPC - 300mm² AAAC equivalent bare": "Upas OPPC 300mm²",
    # ACSR
    "Gopher - 25mm² ACSR bare (1000m drums)": "Gopher 25mm²",
    "Caton - 25mm² Compacted ACSR bare (1000m drums)": "Caton 25mm²",
    "Rabbit - 50mm² ACSR bare (1000m drums)": "Rabbit 50mm²",
    "Wolf - 150mm² ACSR bare (1000m drums)": "Wolf 150mm²",
    "Horse - 70mm² ACSR bare": "Horse 70mm²",
    "Dog - 100mm² ACSR bare (1000m drums)": "Dog 100mm²",
    "Dingo - 150mm² ACSR bare (1000m drums)": "Dingo 150mm²",
    # Copper
    "Hard Drawn Copper 16mm² ( 3/2.65mm ) (500m drums)": "Copper 16mm²",
    "Hard Drawn Copper 32mm² ( 3/3.75mm ) (1000m drums)": "Copper 32mm²",
    "Hard Drawn Copper 70mm² (500m drums)": "Copper 70mm²",
    "Hard Drawn Copper 100mm² (500m drums)": "Copper 100mm²",
    # PVC covered
    "35mm² Copper (Green / Yellow PVC covered) (50m drums)": "Copper 35mm² GY PVC",
    "70mm² Copper (Green / Yellow PVC covered) (50m drums)": "Copper 70mm² GY PVC",
    "35mm² Copper (Blue PVC covered) (50m drums)": "Copper 35mm² Blue PVC",
    "70mm² Copper (Blue PVC covered) (50m drums)": "Copper 70mm² Blue PVC",
    # Double insulated
    "35mm² Double Insulated (Brown) (50m drums)": "Double Insulated 35mm² Brown",
    "35mm² Double Insulated (Blue) (50m drums)": "Double Insulated 35mm² Blue",
    "70mm² Double Insulated (Brown) (50m drums)": "Double Insulated 70mm² Brown",
    "70mm² Double Insulated (Blue) (50m drums)": "Double Insulated 70mm² Blue",
    "120mm² Double Insulated (Brown) (50m drums)": "Double Insulated 120mm² Brown",
    "120mm² Double Insulated (Blue) (50m drums)": "Double Insulated 120mm² Blue"
}

    # LV cables per meter
conductor_2_keys = {
    "LV Cable 1ph 4mm Concentric (250m drums)": "LV 1ph 4mm Concentric",
    "LV Cable 1ph 25mm CNE (250m drums)": "LV 1ph 25mm CNE",
    "LV Cable 1ph 25mm SNE (100m drums)": "LV 1ph 25mm SNE",
    "LV Cable 1ph 35mm CNE (250m drums)": "LV 1ph 35mm CNE",
    "LV Cable 1ph 35mm SNE (100m drums)": "LV 1ph 35mm SNE",
    "LV Cable 3ph 35mm Cu Split Con (250m drums)": "LV 3ph 35mm Cu Split Con",
    "LV Cable 3ph 35mm SNE (250m drums)": "LV 3ph 35mm SNE",
    "LV Cable 3ph 35mm CNE (250m drums)": "LV 3ph 35mm CNE",
    "LV Cable 3ph 35mm CNE Al (LSOH) (250m drums)": "LV 3ph 35mm CNE Al LSOH",
    "LV Cable 3c 95mm W/F (250m drums)": "LV 3c 95mm W/F",
    "LV Cable 3c 185mm W/F (250m drums)": "LV 3c 185mm W/F",
    "LV Cable 3c 300mm W/F (250m drums)": "LV 3c 300mm W/F",
    "LV Cable 4c 95mm W/F (250m drums)": "LV 4c 95mm W/F",
    "LV Cable 4c 185mm W/F (250m drums)": "LV 4c 185mm W/F",
    "LV Cable 4c 240mm W/F (250m drums)": "LV 4c 240mm W/F",
    "LV Marker Tape (365m roll)": "LV Marker Tape",
    # 11kV
    "11kv Cable 95mm 3c Poly (250m drums)": "11kV 3c 95mm Poly",
    "11kv Cable 185mm 3c Poly (250m drums)": "11kV 3c 185mm Poly",
    "11kv Cable 300mm 3c Poly (250m drums)": "11kV 3c 300mm Poly",
    "11kv Cable 95mm 1c Poly (250m drums)": "11kV 1c 95mm Poly",
    "11kv Cable 185mm 1c Poly (250m drums)": "11kV 1c 185mm Poly",
    "11kv Cable 300mm 1c Poly (250m drums)": "11kV 1c 300mm Poly",
    "11kV Marker Tape (40m roll)": "11kV Marker Tape"
}


equipment_keys = {
    "Noja": "Noja",
    "11kV PMSW (Soule)": "11kV PMSW (Soule)",
    "11kv ABSW Hookstick Standard": "11kv ABSW Hookstick Standard",
    "11kv ABSW Hookstick Spring loaded mech": "11kv ABSW Hookstick Spring loaded mech",
    "33kv ABSW Hookstick Dependant": "33kv ABSW Hookstick Dependant",
    "100A LV Fuse JPU 82.5mm": "100A LV Fuse JPU 82.5mm",
    "160A LV Fuse JPU 82.5mm": "160A LV Fuse JPU 82.5mm",
    "200A LV Fuse JPU 82.5mm": "200A LV Fuse JPU 82.5mm",
    "315A LV Fuse JPU 82.5mm": "315A LV Fuse JPU 82.5mm",
    "400A LV Fuse JPU 82.5mm": "400A LV Fuse JPU 82.5mm",
    "200A LV Fuse JSU 92mm": "200A LV Fuse JSU 92mm",
    "315A LV Fuse JSU 92mm": "315A LV Fuse JSU 92mm",
    "400A LV Fuse JSU 92mm": "400A LV Fuse JSU 92mm",
    "100A LV Fuse - Porcelain screw-in": "100A LV Fuse - Porcelain screw-in",
    "160A LV Fuse - Porcelain screw-in": "160A LV Fuse - Porcelain screw-in",
    "200A LV Fuse - Porcelain screw-in": "200A LV Fuse - Porcelain screw-in",
    "Single Phase cut out kit 100A Henley Series 7": "Single Phase cut out kit 100A Henley Series 7",
    "Single Phase SNE Sealing Chamber": "Single Phase SNE Sealing Chamber",
    "Three Phase cut out kit 100A Henley Series 7": "Three Phase cut out kit 100A Henley Series 7",
    "Three Phase 200A Cut out": "Three Phase 200A Cut out",
    "Earth Connector Block 100A 5 Way": "Earth Connector Block 100A 5 Way",
    "Cut out Fuse (MF) 60A": "Cut out Fuse (MF) 60A",
    "Cut out Fuse (MF) 80A": "Cut out Fuse (MF) 80A",
    "Cut out Fuse (MF) 100A": "Cut out Fuse (MF) 100A",
    "Temporary Meter seal white plastic (100)": "Temporary Meter seal white plastic (100)",
    "Meter seals for use with sealing pliers (100)": "Meter seals for use with sealing pliers (100)",
    "Meter sealing wire 200mm long (each)": "Meter sealing wire 200mm long (each)",
    "ABC 1PH & 3PH TERM BOX": "ABC 1PH & 3PH TERM BOX",
    "SINGLE PHASE FUSED ABC BOX": "SINGLE PHASE FUSED ABC BOX",
    "1PH & 3PH FUSED SERV WALL BOX": "1PH & 3PH FUSED SERV WALL BOX",
    "25mm Galvanised Conduit": "25mm Galvanised Conduit",
    "25mm Galvanised Conduit saddles": "25mm Galvanised Conduit saddles",
    "Street Lighting Cut out CNE": "Street Lighting Cut out CNE",
    "Street Lighting Cut out SNE": "Street Lighting Cut out SNE",
    "11KV FUSE UNIT - C-TYPE": "11KV FUSE UNIT - C-TYPE",
    "11KV SOLID LINK - C-TYPE": "11KV SOLID LINK - C-TYPE",
    "11KV OHL ASL C-TYPE RESET 20A 2 SHOT": "11KV OHL ASL C-TYPE RESET 20A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 25A 2 SHOT": "11KV OHL ASL C-TYPE RESET 25A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 40A 1 SHOT": "11KV OHL ASL C-TYPE RESET 40A 1 SHOT",
    "11KV OHL ASL C-TYPE RESET 40A 2 SHOT": "11KV OHL ASL C-TYPE RESET 40A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 63A 1 SHOT": "11KV OHL ASL C-TYPE RESET 63A 1 SHOT",
    "11KV OHL ASL C-TYPE RESET 63A 2 SHOT": "11KV OHL ASL C-TYPE RESET 63A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 63A 3 SHOT": "11KV OHL ASL C-TYPE RESET 63A 3 SHOT",
    "11KV OHL ASL C-TYPE RESET 100A 1 SHOT": "11KV OHL ASL C-TYPE RESET 100A 1 SHOT",
    "11KV OHL ASL C-TYPE RESET 100A 2 SHOT": "11KV OHL ASL C-TYPE RESET 100A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 100A 3 SHOT": "11KV OHL ASL C-TYPE RESET 100A 3 SHOT",
    "11KV FUSE CARRIER - C-TYPE": "11KV FUSE CARRIER - C-TYPE",
    "11KV OHL FUSE ELEMENT C-TYPE 15A": "11KV OHL FUSE ELEMENT C-TYPE 15A",
    "11KV OHL FUSE ELEMENT C-TYPE 25A": "11KV OHL FUSE ELEMENT C-TYPE 25A",
    "11KV OHL FUSE ELEMENT C-TYPE 30A": "11KV OHL FUSE ELEMENT C-TYPE 30A",
    "11KV OHL FUSE ELEMENT C-TYPE 40A": "11KV OHL FUSE ELEMENT C-TYPE 40A",
    "11KV OHL FUSE ELEMENT C-TYPE 50A": "11KV OHL FUSE ELEMENT C-TYPE 50A",
    "11KV OHL ASL - CHEMICAL ACTUATOR": "11KV OHL ASL - CHEMICAL ACTUATOR",
    "11KV OHL ASL DJP-TYPE 20A 2 SHOT": "11KV OHL ASL DJP-TYPE 20A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 25A 1 SHOT": "11KV OHL ASL DJP-TYPE 25A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 25A 2 SHOT": "11KV OHL ASL DJP-TYPE 25A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 40A 1 SHOT": "11KV OHL ASL DJP-TYPE 40A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 40A 2 SHOT": "11KV OHL ASL DJP-TYPE 40A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 63A 1 SHOT": "11KV OHL ASL DJP-TYPE 63A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 63A 2 SHOT": "11KV OHL ASL DJP-TYPE 63A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 63A 3 SHOT": "11KV OHL ASL DJP-TYPE 63A 3 SHOT",
    "11KV OHL ASL DJP-TYPE 100A 1 SHOT": "11KV OHL ASL DJP-TYPE 100A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 100A 2 SHOT": "11KV OHL ASL DJP-TYPE 100A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 100A 3 SHOT": "11KV OHL ASL DJP-TYPE 100A 3 SHOT",
    "11KV OHL FUSE ELEMENT DJP-TYPE 15A": "11KV OHL FUSE ELEMENT DJP-TYPE 15A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 25A": "11KV OHL FUSE ELEMENT DJP-TYPE 25A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 30A": "11KV OHL FUSE ELEMENT DJP-TYPE 30A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 40A": "11KV OHL FUSE ELEMENT DJP-TYPE 40A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 50A": "11KV OHL FUSE ELEMENT DJP-TYPE 50A",
    "0.5 kVa Tx for Noja": "0.5 kVa Tx for Noja",
    "Military Cable for Noja": "Military Cable for Noja",
    "Antenna for Soule or Noja": "Antenna for Soule or Noja",
    "Bracket for antenna": "Bracket for antenna",
    "Coax cable (5m)": "Coax cable (5m)",
    "Antenna for Soule or Noja": "Antenna for Soule or Noja",
    "Bracket for antenna": "Bracket for antenna",
    "Coax cable (5m)": "Coax cable (5m)",
}

insulator_keys = {
    "11kV Pin Insulator; Polymeric": "11kV Pin Insulator; Polymeric",
    "11kV Pin Insulator; Polymeric; High Creepage": "11kV Pin Insulator; Polymeric; High Creepage",
    "33kV Pin Insulator; Porcelain": "33kV Pin Insulator; Porcelain",
    "33kV Post Insulator; Polymeric; Clamp Top Plate": "33kV Post Insulator; Polymeric; Clamp Top Plate",
    "36kV Composite Post Groove Top": "36kV Composite Post Groove Top",
    "11kV Tension Insulator; Polymeric (70kN)": "11kV Tension Insulator; Polymeric (70kN)",
    "33kV Tension Insulator; Polymeric (70kN)": "33kV Tension Insulator; Polymeric (70kN)",
    "36kV Composite Tension Ball/Socket Fitting (125 kN)": "36kV Composite Tension Ball/Socket Fitting (125 kN)",
    "LV / 11kV Stay Insulator": "LV / 11kV Stay Insulator",
    "33kV Stay Insulator": "33kV Stay Insulator",
    "LV Insulator Bobbin Type": "LV Insulator Bobbin Type",
    "LV Insulator Coachscrew Type": "LV Insulator Coachscrew Type"
}


lv_joint_kit_keys = {
    "LVKIT/001": "LVKIT/001 Straight Jt Kit 35mm 1ph CNE/SNE Plastic",
    "LVKIT/002": "LVKIT/002 Straight Jt Kit 35mm 1ph CNE/SNE Pilc",
    "LVKIT/003": "LVKIT/003 Straight Jt Kit 35mm 3ph CNE/SNE Plastic",
    "LVKIT/004": "LVKIT/004 Staight Jt 3ph 35mm XLPE to 4-35 PILC",
    "LVKIT/005": "LVKIT/005 LV Service Cable Stop End",
    "LVKIT/006": "LVKIT/006 LV Service off a service 4-35mm 1/3 phase CNE/SNE",
    "LVKIT/007": "LVKIT/007 LV Service off a service 4-35mm PILC 1ph CNE/SNE",
    "LVKIT/008": "LVKIT/008 Service Pole Term to OHL 1PH CNE",
    "LVKIT/009": "LVKIT/009 Service Pole Term to OHL 1PH SNE",
    "LVKIT/010": "LVKIT/010 Service Pole Term to OHL 3PH 35mm",
    "LVKIT/011": "LVKIT/011 Service Pole Term to Fuses 1PH CNE",
    "LVKIT/012": "LVKIT/012 Service Pole Term to Fuses 1PH SNE",
    "LVKIT/013": "LVKIT/013 Service Pole Term to Fuses 3PH 35mm",
    "LVKIT/014": "LVKIT/014 Service Breech Joint 70-185mm 3c W/F - CNE/SNE",
    "LVKIT/015": "LVKIT/015 Service Breech Joint 240-300mm 3c W/F - CNE/SNE",
    "LVKIT/016": "LVKIT/016 Service Breech Joint 50-95mm PILC - CNE/SNE",
    "LVKIT/017": "LVKIT/017 Service Breech Joint 95-185mm PILC - CNE/SNE",
    "LVKIT/018": "LVKIT/018 Service Breech Joint 185-300mm PILC - CNE/SNE",
    "LVKIT/019": "LVKIT/019 Straight Joint up to 95mm 3c W/F / PILC",
    "LVKIT/020": "LVKIT/020 Straight Joint 185mm 3c W/F / PILC / CONSAC",
    "LVKIT/021": "LVKIT/021 Straight Joint 300mm 3c W/F / PILC / CONSAC",
    "LVKIT/022": "LVKIT/022 Mains Breech Joint 70-95mm 3c W/F",
    "LVKIT/023": "LVKIT/023 Mains Breech Joint 185mm 3c W/F",
    "LVKIT/024": "LVKIT/024 Mains Breech Joint 240/300mm 3c W/F",
    "LVKIT/025": "LVKIT/025 Mains Breech Joint 70-95mm W/F / 50-95mm PILC",
    "LVKIT/026": "LVKIT/026 Mains Breech Joint 185mm W/F / 95-185mm PILC",
    "LVKIT/027": "LVKIT/027 Mains Breech Joint 240/300mm W/F / 185-300mm PILC",
    "LVKIT/028": "LVKIT/028 Loop / V Joint 50-95mm W/F / PILC",
    "LVKIT/029": "LVKIT/029 Loop / V Joint >95-300mm W/F / PILC",
    "LVKIT/030": "LVKIT/030 Y / 3 Loose end Joint 50-185mm W/F / PILC / Districable",
    "LVKIT/031": "LVKIT/031 Y / 3 Loose end Joint 185-300mm W/F / PILC / Districable",
    "LVKIT/032": "LVKIT/032 Stop End 70-95mm W/F / CONSAC",
    "LVKIT/033": "LVKIT/033 Stop End 185-300mm W/F / CONSAC",
    "LVKIT/034": "LVKIT/034 Stop End 50-95mm PILC",
    "LVKIT/035": "LVKIT/035 Stop End 95-300mm PILC",
    "LVKIT/037": "LVKIT/037 Pole Term to OHL 70-95mm W/F",
    "LVKIT/038": "LVKIT/038 Pole Term to OHL 185mm W/F",
    "LVKIT/039": "LVKIT/039 Pole Term to Fuses 70-95mm W/F",
    "LVKIT/040": "LVKIT/040 Pole Term to Fuses 185mm W/F"
}


lv_joint_module_keys = {
    "LVMOD/001": "LVMOD/001 Armour bond module for PILC Service cable Stop Ends",
    "LVMOD/002": "LVMOD/002 Branch connector module for service cables",
    "LVMOD/003": "LVMOD/003 Phase connector remake module for service cables",
    "LVMOD/004": "LVMOD/004 XL Brass tunnel connector module for old PILC concentric cables",
    "LVMOD/005": "LVMOD/005 Insulated insulating piercing mains/service branch connector module (up to 185mm2)",
    "LVMOD/006": "LVMOD/006 Insulated insulating piercing mains/service branch connector module (240-300mm2)",
    "LVMOD/007": "LVMOD/007 Brass neutral earth connector module",
    "LVMOD/008": "LVMOD/008 CONSAC Brass neutral earth connector module",
    "LVMOD/009": "LVMOD/009 95mm2 straight type channel connector module",
    "LVMOD/011": "LVMOD/011 185mm2 straight type channel connector module",
    "LVMOD/013": "LVMOD/013 300mm2 straight type channel connector module",
    "LVMOD/015": "LVMOD/015 95mm2 branch type channel connector module",
    "LVMOD/017": "LVMOD/017 185mm2 branch type channel connector module",
    "LVMOD/018": "LVMOD/018 185mm2 branch type channel connector c/w brass non-shear bolts module",
    "LVMOD/019": "LVMOD/019 300mm2 branch type channel connector module",
    "LVMOD/021": "LVMOD/021 95mm2 1/2 length branch type connector module",
    "LVMOD/022": "LVMOD/022 300mm2 1/2 length branch type connector module",
    "LVMOD/023": "LVMOD/023 95mm2 Service Bridge Piece module",
    "LVMOD/024": "LVMOD/024 185mm2 Service Bridge Piece module",
    "LVMOD/025": "LVMOD/025 300mm2 Service Bridge Piece module",
    "LVMOD/026": "LVMOD/026 upto 35mm2 PILC service cable Earth Bond Kit module",
    "LVMOD/027": "LVMOD/027 50-95mm2 PILC Mains cable Earth Bond Kit module",
    "LVMOD/028": "LVMOD/028 >95-185mm2 PILC Mains cable Earth Bond Kit module",
    "LVMOD/029": "LVMOD/029 >185-300mm2 PILC Mains cable Earth Bond Kit module",
    "LVMOD/030": "LVMOD/030 Torque Limiting shear-off device module",
    "LVMOD/031": "LVMOD/031 95mm2 Aluminium mechanical shear-off lug module",
    "LVMOD/032": "LVMOD/032 185mm2 Aluminium mechanical shear-off lug module",
    "LVMOD/033": "LVMOD/033 300mm2 Aluminium mechanical shear-off lug module",
    "LVMOD/034": "LVMOD/034 480-740mm2 range taking Aluminium mechanical shear-off lug module",
    "LVMOD/035": "LVMOD/035 95mm2 Aluminium mechanical shear-off Busbar connector module",
    "LVMOD/036": "LVMOD/036 185mm2 Aluminium mechanical shear-off Busbar connector module",
    "LVMOD/037": "LVMOD/037 300mm2 Aluminium mechanical shear-off Busbar connector module",
    "LVMOD/038": "LVMOD/038 70-95mm2 pole termination module kit for 4c overhead lines and fuses",
    "LVMOD/039": "LVMOD/039 185mm pole termination module kit for 4c overhead lines and fuses",
    "LVMOD/040": "LVMOD/040 35-70mm2 Brass shear off lug module",
    "LVMOD/041": "LVMOD/041 60-120mm2 Brass shear off lug module"
}

hv_joint_termination_keys = {
    "11kv XLPE 3c Straight joint": "11kV XLPE 3c Straight Joint",
    "11kV 95mm XLPE trif joint": "11kV 95mm XLPE Trifurcating Joint",
    "11kV 185 - 300mm XLPE Trif joint": "11kV 185-300mm XLPE Trifurcating Joint",
    "11kV up to 70mm PILC/PICAS to XLPE Joint": "11kV PILC/PICAS to XLPE Joint (up to 70mm)",
    "11kV 95-185 PILC/PICAS to XLPE Joint": "11kV PILC/PICAS to XLPE Joint (95-185mm)",
    "11kV 185-300 PILC/PICAS to XLPE Joint": "11kV PILC/PICAS to XLPE Joint (185-300mm)",
    "11kV 95-185 XLPE to up to 70mm PILC/PICAS Transition Trif Joint": "11kV XLPE to PILC/PICAS Transition Trif Joint (95-185mm to 70mm)",
    "11kV 95-185 XLPE to 95-185 PILC/PICAS Transition Trif Joint": "11kV XLPE to PILC/PICAS Transition Trif Joint (95-185mm)",
    "11kV 185-300 XLPE to 185-300 PILC/PICAS Transition Trif Joint": "11kV XLPE to PILC/PICAS Transition Trif Joint (185-300mm)",
    "11kV Earthing kit for CORAL cables": "11kV Earthing Kit for CORAL Cables",
    "11kV Earthing kit for 50-300mm PILC cables": "11kV Earthing Kit for PILC Cables (50-300mm)",
    "11kV Earthing kit for up to 50mm PILC cables": "11kV Earthing Kit for PILC Cables (up to 50mm)",
    "11kV Build up kit for PILC / CORAL cables": "11kV Build Up Kit for PILC/CORAL Cables",
    "11kV Build up kit for XLPE cables": "11kV Build Up Kit for XLPE Cables",
    "11kV 95/185mm module for PAPER to PAPER joint": "11kV Paper to Paper Joint Module (95/185mm)",
    "11kV 300mm module for PAPER to PAPER joint": "11kV Paper to Paper Joint Module (300mm)",
    "11kV pole Term 1c 95mm": "11kV Pole Termination 1c 95mm",
    "11kV pole Term 1c 185/300mm": "11kV Pole Termination 1c 185/300mm",
    "11kV pole Term 3c 95mm": "11kV Pole Termination 3c 95mm",
    "11kV pole Term 3c 185/300mm": "11kV Pole Termination 3c 185/300mm",
    "OUTDR TERMN POLE STEELWORK 11 KV": "11kV Outdoor Pole Termination Steelwork",
    "11kV 95mm cable clamp for crucifix": "11kV Cable Clamp for Crucifix (95mm)",
    "11kV 185mm cable clamp for crucifix": "11kV Cable Clamp for Crucifix (185mm)",
    "11kV Surge Arrestor (Each)": "11kV Surge Arrestor",
    "33kv Joint Transition Trif (H-Type)": "33kV Joint Transition Trifurcating (H-Type)",
    "33kv Joint Trif (HSL-Type)": "33kV Joint Trifurcating (HSL-Type)",
    "33kv Joint 0.1 sq inch connectors (3 phases)": "33kV Joint Connectors 0.1 sq inch",
    "33kv Joint 0.4/0.5 sq inch connector (per phase)": "33kV Joint Connector 0.4/0.5 sq inch",
    "33kv Joint Connectors for Trif 150/300 Pilc": "33kV Joint Connectors for Trifurcating 150/300 PILC",
    "33kv Joint Straight up to 240mm (per phase)": "33kV Straight Joint (up to 240mm)",
    "33kv Joint Straight over 240mm needs connector (per phase)": "33kV Straight Joint (over 240mm)",
    "33kv Joint 400mm connector (each)": "33kV Joint Connector 400mm",
    "33kv Joint Transition 150/240mm to 0.3 PILC (per phase)": "33kV Joint Transition 150/240mm to 0.3 PILC",
    "11/33kv Pot End module up to 300mm (3 phases)": "11/33kV Pot End Module (up to 300mm)",
    "33kV Pole Term 1c 150-240mm (3 phase set)": "33kV Pole Termination 1c 150-240mm",
    "33kV Pole Term 1c 400-630mm (3 phase set)": "33kV Pole Termination 1c 400-630mm",
    "33kV Cable cleats for pole terms": "33kV Cable Cleats for Pole Terminations",
    "33kV Surge Arrestor 36kV (Each)": "33kV Surge Arrestor 36kV"
}

cable_accessory_keys = {
    "End cap up to 17mm diameter (25(1))": "End cap up to 17mm diameter (25(1))",
    "End cap 17-30mm dia(35(3))": "End cap 17-30mm dia(35(3))",
    "End Cap 30-45mm dia (95 LV or HV)": "End Cap 30-45mm dia (95 LV or HV)",
    "End Cap 45-95mm dia (185-300 LV or HV)": "End Cap 45-95mm dia (185-300 LV or HV)",
    "Ducting 32mm (OD 38mm) per metre (100m coil)": "Ducting 32mm (OD 38mm) per metre (100m coil)",
    "Ducting 50mm (OD 58mm) per metre (50m coil)": "Ducting 50mm (OD 58mm) per metre (50m coil)",
    "Ducting 100mm (3m Length) (90 in pallet)": "Ducting 100mm (3m Length) (90 in pallet)",
    "Ducting bend (100mm / 11.25 degree)": "Ducting bend (100mm / 11.25 degree)",
    "Ducting bend (100mm / 22.5 degree)": "Ducting bend (100mm / 22.5 degree)",
    "Ducting bend (100mm / 45 degree)": "Ducting bend (100mm / 45 degree)",
    "Ducting 150mm (3m Length) (39 in pallet)": "Ducting 150mm (3m Length) (39 in pallet)",
    "Ducting bend (150mm / 11.25 degree)": "Ducting bend (150mm / 11.25 degree)",
    "Ducting bend (150mm / 22.5 degree)": "Ducting bend (150mm / 22.5 degree)",
    "Ducting bend (150mm / 45 degree)": "Ducting bend (150mm / 45 degree)",
    "Resin 2 litre JEM Permanent": "Resin 2 litre JEM Permanent",
    "Resin 6 litre JEM Permanent": "Resin 6 litre JEM Permanent"
}

foundation_steelwork_keys = {
    "H' Pole Foundation Brace Steelwork for P6.010mm Centres ( Ref. SP4017651 )": "H' Pole Foundation Brace Steelwork for P6.010mm Centres ( Ref. SP4017651 )",
    "'H' Pole Foundation Brace Steelwork for 2500mm Centres ( Ref. SP4017652 )": "'H' Pole Foundation Brace Steelwork for 2500mm Centres ( Ref. SP4017652 )",
    "Stay / Foundation Block Type 1; 850mm as SP4019020": "Stay / Foundation Block Type 1; 850mm as SP4019020",
    "Stay / Foundation Block Type 2; 1300mm as SP4019020": "Stay / Foundation Block Type 2; 1300mm as SP4019020",
    "Foundation Block Type 3; 1500mm as SP4019020": "Foundation Block Type 3; 1500mm as SP4019020"
}

summary_items = [
    "Erect Single HV/EHV Pole, up to and including 12 metre pole.",
    "Erect Section Structure 'H' HV/EHV Pole, up to and including 12 metre pole",
    "Erect LV Structure Single Pole, up to and including 12 metre pole",
    "Recover single pole, up to and including 15 metres in height, and reinstate, all ground conditions",
    "Recover 'A' / 'H' pole, up to and including 15 metres in height, and reinstate, all ground conditions",
    "Erect 11kV/33kV ABSW.",
    "Erect 11kV Remote Controlled Switch Disconnector ( Soule Auguste ) or Auto Reclosure unit c/w VT, Aerial, RTU & umbilical cable",
    "Erect pole mounted transformer up to 100kVA 1.ph",
    "Erect pole mounted transformer up to 200kVA 3.p.h",
    "Remove pole mounted transformer",
    "Remove platform mounted or 'H' pole mounted transformer",
    "Remove 11kV/33kV ABSW",
    "Remove Auto Reclosure",
    "Install bare conductor, run out, sag, terminate, bind in and connect jumpers; <100mm²",
    "Install bare conductor, run out, sag, terminate, bind in and connect jumpers; >=100mm² <200mm²",
    "Install conductor, run out, sag, terminate, clamp in and connect jumpers; 2c + Earth",
    "Install conductor, run out, sag, terminate, clamp in and connect jumpers; 4c + Earth",
    "Install service span including connection to mainline & building / structure",
    "Remove 1.ph or 3.ph HV fuses",
    "Erect 3.ph fuse units at single tee off pole or in line pole"
    
]

categories = [
    ("Poles 🪵", pole_keys, "Quantity"),
    ("Poles _erected 🪵", pole_erected_keys, "Quantity"),
    ("Poles _replaced 🪵", poles_replaced_keys, "Quantity"),
    ("Transformers ⚡🏭", transformer_keys, "Quantity"),
    ("Conductors", conductor_keys, "Length (Km)"),
    ("Conductors_2", conductor_2_keys, "Length (Km)"),
    ("Equipment", equipment_keys, "Quantity"),
    ("Insulators", insulator_keys, "Quantity"),
    ("LV Joints (Kits)", lv_joint_kit_keys, "Quantity"),
    ("LV Joint Modules", lv_joint_module_keys, "Quantity"),
    ("HV Joints / Terminations ⚡", hv_joint_termination_keys, "Quantity"),
    ("Cable Accessories 🔌", cable_accessory_keys, "Quantity"),
    ("Foundation & Steelwork 🏗️", foundation_steelwork_keys, "Quantity")
]

column_rename_map = {
    "mapped": "Output",
    "segmentcode": "Circuit",
    "datetouse_display": "Date",
    "qty": "Quantity_original",
    "qsub": "Quantity_used",
    "segmentdesc": "Segment",
    "shire": "District",
    "pid_ohl_nr": "PID",
    "projectmanager": "Project Manager"
}

export_columns = [
    'Output','comment', 'item', 'Quantity_original','Quantity_used', 'material_code','type', 'pole', 'Date',
    'District', 'project', 'Project Manager', 'Circuit', 'Segment',
    'team lider', 'PID', 'sourcefile'
]

# --- Gradient background ---
gradient_bg = """
<style>
    .stApp {
        background: linear-gradient(
            90deg,
            rgba(41, 28, 66, 1) 10%, 
            rgba(36, 57, 87, 1) 35%
        );
        color: white;
    }
</style>
"""
st.markdown(gradient_bg, unsafe_allow_html=True)

# --- Load logos ---
logo_left = Image.open(r"Images/GaeltecImage.png").resize((80, 80))
logo_right = Image.open(r"Images/SPEN.png").resize((160, 80))

# --- Header layout ---
col1, col2, col3 = st.columns([1, 4, 1])
with col1: st.image(logo_left)
with col2: st.markdown("<h1 style='text-align:center; margin:0;'>Gaeltec Utilities.UK</h1>", unsafe_allow_html=True)
with col3: st.image(logo_right)
st.markdown("<h1>📊 Data Management Dashboard</h1>", unsafe_allow_html=True)

# -------------------------------
# --- File Upload & Initial DF ---
# -------------------------------
# --- Upload Aggregated Parquet file ---
# --- Load aggregated Parquet file ---
# -------------------------------
# App Header
# -------------------------------
st.header("Upload Data Files")

# -------------------------------
# Load Aggregated Parquet
# -------------------------------
master_file = st.file_uploader(
    "Upload Master.parquet",
    type=["parquet"],
    key="master"
)

resume_file = st.file_uploader(
    "Upload CF_resume.parquet",
    type=["parquet"],
    key="resume_file"
)
resume_df = None

if resume_file is not None:
    resume_df = pd.read_parquet(resume_file)
    resume_df.columns = resume_df.columns.str.strip().str.lower()

misc_file = st.file_uploader(
    "Upload miscellaneous.parquet",
    type=["parquet"],
    key="misc_file"
)
misc_df = None

if misc_file is not None:
    try:
        misc_df = pd.read_parquet(misc_file)
        misc_df.columns = misc_df.columns.str.strip().str.lower()
    except Exception as e:
        st.warning(f"Could not load Miscellaneous parquet: {e}")

base_df = None
st.header("Upload Data Files")

agg_view = None

if master_file is not None:
    df = pd.read_parquet(master_file)
    df.columns = df.columns.str.strip().str.lower()  # normalize columns

    if 'datetouse' in df.columns:
        df['datetouse_dt'] = pd.to_datetime(df['datetouse'], errors='coerce')
        df['datetouse_display'] = df['datetouse_dt'].dt.strftime("%d/%m/%Y")
        df.loc[df['datetouse_dt'].isna(), 'datetouse_display'] = "Unplanned"
        df['datetouse_dt'] = df['datetouse_dt'].dt.normalize()
    else:
        df['datetouse_dt'] = pd.NaT
        df['datetouse_display'] = "Unplanned"

    agg_view = df.copy()

# -------------------------------
# --- Team Filter (GLOBAL) ---
# -------------------------------
base_df = None

if master_file:
    base_df = pd.read_parquet(master_file)
    base_df.columns = base_df.columns.str.strip().str.lower()

    # Normalize date
    if 'datetouse' in base_df.columns:
        base_df['datetouse_dt'] = pd.to_datetime(base_df['datetouse'], errors='coerce').dt.normalize()
    else:
        base_df['datetouse_dt'] = pd.NaT

    # Normalize numeric columns
    for col in ['total', 'orig']:
        if col in base_df.columns:
            base_df[col] = (
                base_df[col]
                .astype(str)
                .str.replace(" ", "")
                .str.replace(",", ".", regex=False)
            )
            base_df[col] = pd.to_numeric(base_df[col], errors='coerce')

# Stop early if no data
if base_df is None:
    st.info("Please upload Master.parquet to continue.")
    st.stop()

# -------------------------------
# Sidebar Filters
# -------------------------------
st.sidebar.header("Filter Options")

def multiselect_filter(df, column, label):
    if column not in df.columns:
        return ["All"], df
    options = ["All"] + sorted(df[column].dropna().astype(str).unique())
    selected = st.sidebar.multiselect(label, options, default=["All"])
    if "All" not in selected:
        df = df[df[column].astype(str).isin(selected)]
    return selected, df

filtered_df = base_df.copy()

selected_shire, filtered_df = multiselect_filter(filtered_df, 'shire', "Select Shire")
selected_project, filtered_df = multiselect_filter(filtered_df, 'project', "Select Project")
selected_pm, filtered_df = multiselect_filter(filtered_df, 'projectmanager', "Select Project Manager")
selected_segment, filtered_df = multiselect_filter(filtered_df, 'segmentcode', "Select Segment Code")
selected_pole, filtered_df = multiselect_filter(filtered_df, 'pole', "Select Pole")
selected_type, filtered_df = multiselect_filter(filtered_df, 'type', "Select Type")
selected_team, filtered_df = multiselect_filter(filtered_df, 'team_name', "Select Team")


# -------------------------------
# Date Filter
# -------------------------------
filter_type = st.sidebar.selectbox(
    "Filter by Date",
    ["Single Day", "Week", "Month", "Year", "Custom Range", "Unplanned"]
)

date_range_str = ""
filtered_df['datetouse_dt'] = pd.to_datetime(filtered_df['datetouse_dt'])

if filter_type == "Unplanned":
    filtered_df = filtered_df[filtered_df['datetouse_dt'].isna()]
    date_range_str = "Unplanned"

else:
    filtered_df = filtered_df[filtered_df['datetouse_dt'].notna()]

    if filter_type == "Single Day":
        d = st.sidebar.date_input("Select date")
        filtered_df = filtered_df[filtered_df['datetouse_dt'] == pd.Timestamp(d)]
        date_range_str = str(d)

    elif filter_type == "Week":
        start = pd.Timestamp(st.sidebar.date_input("Week start"))
        end = start + pd.Timedelta(days=6)
        filtered_df = filtered_df[
            (filtered_df['datetouse_dt'] >= start) &
            (filtered_df['datetouse_dt'] <= end)
        ]
        date_range_str = f"{start} → {end}"

    elif filter_type == "Month":
        d = st.sidebar.date_input("Pick any date in month")
        filtered_df = filtered_df[
            (filtered_df['datetouse_dt'].dt.month == d.month) &
            (filtered_df['datetouse_dt'].dt.year == d.year)
        ]
        date_range_str = d.strftime("%B %Y")

    elif filter_type == "Year":
        y = st.sidebar.number_input("Year", 2000, 2100, 2025)
        filtered_df = filtered_df[filtered_df['datetouse_dt'].dt.year == y]
        date_range_str = str(y)

    elif filter_type == "Custom Range":
        start = pd.Timestamp(st.sidebar.date_input("Start date"))
        end = pd.Timestamp(st.sidebar.date_input("End date"))
        filtered_df = filtered_df[
            (filtered_df['datetouse_dt'] >= start) &
            (filtered_df['datetouse_dt'] <= end)
        ]
        date_range_str = f"{start} → {end}"

    # -------------------------------
    # --- Total & Variation Display ---
    # -------------------------------
    total_sum, variation_sum = 0, 0
    if 'total' in filtered_df.columns:
        total_series = pd.to_numeric(filtered_df['total'].astype(str).str.replace(" ", "").str.replace(",", ".", regex=False),
                                     errors='coerce')
        total_sum = total_series.sum(skipna=True)
        if 'orig' in filtered_df.columns:
            orig_series = pd.to_numeric(filtered_df['orig'].astype(str).str.replace(" ", "").str.replace(",", ".", regex=False),
                                        errors='coerce')
            variation_sum = (total_series - orig_series).sum(skipna=True)

    formatted_total = f"{total_sum:,.2f}".replace(",", " ").replace(".", ",")
    formatted_variation = f"{variation_sum:,.2f}".replace(",", " ").replace(".", ",")

    # Money logo
    money_logo_path = r"Images/Pound.png"
    money_logo = Image.open(money_logo_path).resize((40, 40))
    buffered = BytesIO()
    money_logo.save(buffered, format="PNG")
    money_logo_base64 = base64.b64encode(buffered.getvalue()).decode()

    # Display Total & Variation (Centered)
    st.markdown("<h2>Financial</h2>", unsafe_allow_html=True)
    st.markdown("<h3 style='text-align:center; color:white;'>Revenue</h3>", unsafe_allow_html=True)
    try:
        st.markdown(
            f"""
            <div style='display:flex; justify-content:center;'>
                <div style='display:flex; flex-direction:column; gap:4px;'>
                    <div style='display:flex; align-items:center; gap:10px;'>
                        <h2 style='color:#32CD32; margin:0; font-size:36px;'><b>Total:</b> {formatted_total}</h2>
                        <img src='data:image/png;base64,{money_logo_base64}' width='40' height='40'/>
                    </div>
                    <div style='display:flex; align-items:center; gap:8px;'>
                        <h2 style='color:#32CD32; font-size:25px; margin:0;'><b>Variation:</b> {formatted_variation}</h2>
                        <img src='data:image/png;base64,{money_logo_base64}' width='28' height='28'/>
                    </div>
                    <p style='text-align:center; font-size:14px; margin-top:4px;'>
                        ({date_range_str}, Shires: {selected_shire}, Projects: {selected_project}, PMs: {selected_pm})
                    </p>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    except Exception as e:
        st.warning(f"Could not display Total & Variation: {e}")
# -------------------------------
# Revenue Over Time
# -------------------------------

general_summary = pd.DataFrame(
    columns=["Description", "Total Quantity", "Comment"]
)
if not filtered_df.empty and 'datetouse_dt' in filtered_df.columns and 'total' in filtered_df.columns:
    # Aggregate revenue per date
    revenue_df = (
        filtered_df
        .dropna(subset=['datetouse_dt'])
        .groupby('datetouse_dt', as_index=False)['total']
        .sum()
        .sort_values('datetouse_dt')
    )

    # Ensure datetime column
    revenue_df['datetouse_dt'] = pd.to_datetime(revenue_df['datetouse_dt'])

    import plotly.graph_objects as go
    fig = go.Figure()

    # Scatter points (all data)
    fig.add_trace(go.Scattergl(
        x=revenue_df['datetouse_dt'],
        y=revenue_df['total'],
        mode='markers',
        marker=dict(size=8, color='#FFA500'),
        name='Revenue'
    ))

    # Dashed line connecting points
    fig.add_trace(go.Scatter(
        x=revenue_df['datetouse_dt'],
        y=revenue_df['total'],
        mode='lines',
        line=dict(dash='dash', color='#FFA500'),
        name='Trend'
    ))

    # Layout with horizontal gridlines
    fig.update_layout(
        height=500,
        xaxis_title="Date",
        yaxis_title="Revenue (£)",
        hovermode="x unified",
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font=dict(color='white'),
        xaxis=dict(showgrid=True, gridcolor='rgba(255,255,255,0.1)'),
        yaxis=dict(showgrid=True, gridcolor='rgba(255,255,255,0.2)', zeroline=False)
    )

    st.plotly_chart(fig, use_container_width=True)
else:
    st.info("No data for selected filters.")

if filtered_df is not None and not filtered_df.empty:
    buffer_agg = BytesIO()

    with pd.ExcelWriter(buffer_agg, engine="openpyxl") as writer:

        # ---- Prepare export_df ----
        export_df = filtered_df.copy()
        export_df = export_df.rename(columns=column_rename_map)

        if "done" in export_df.columns:
            export_df["done"] = pd.to_datetime(export_df["done"], errors="coerce")
            export_df["done_display"] = export_df["done"].dt.strftime("%d/%m/%Y")
            export_df.loc[export_df["done"].isna(), "done"] = "Unplanned"

        cols_to_include = [
            "item","comment", "Quantity_original", "Quantity_used", "material_code",
            "type", "pole", "datetouse_dt","done_display", "District", "project",
            "Project Manager", "Circuit", "Segment",
            "team lider", "PID", "sourcefile"
        ]
        cols_to_include = [c for c in cols_to_include if c in export_df.columns]
        export_df = export_df[cols_to_include]

        # ---- Output sheet (start below images) ----
        export_df.to_excel(writer, sheet_name="Output", index=False, startrow=1)
        ws = writer.book["Output"]

        # ---- Summary sheet ----
        if "Quantity_used" in export_df.columns:
            # Ensure numeric type
            # Apply normalization
            export_df["Quantity_used"] = pd.to_numeric(export_df["Quantity_used"], errors="coerce").fillna(0)
            special_item = (
                "Erect 11kV Remote Controlled Switch Disconnector (Soule Auguste) or Auto Reclosure unit c/w VT, Aerial, RTU & umbilical cable."
            )
            export_df["item_norm"] = export_df["item"].apply(normalize_item)
            summary_items_norm = [normalize_item(i) for i in summary_items]
            special_item_norm = normalize_item(special_item)
                # Add comments column for the special item
            # Aggregate sum by item
            summary_df = (
                export_df[export_df["item_norm"].isin(summary_items_norm)]
                .groupby("item_norm", as_index=False)["Quantity_used"]
                .sum()
            )

            if not summary_df.empty:
                general_summary = (summary_df.merge(export_df[["item_norm", "item"]],on="item_norm",how="left").drop_duplicates("item_norm")
                                   .rename(columns={"item": "Description","Quantity_used": "Total Quantity"})[["Description", "Total Quantity"]])

                # Ensure Comment column exists
                general_summary["Comment"] = ""

            # Extract all rows for the special item
            special_df = export_df[export_df["item_norm"].str.contains(special_item_norm, na=False)].copy()

            if not special_df.empty:
                # Group by unique comment and sum quantities
                special_summary = (
                    special_df.groupby(["item", "comment"], as_index=False)["Quantity_used"]
                    .sum()
                    .rename(columns={"item": "Description", "Quantity_used": "Total Quantity", "comment": "Comment"})
                    )

                # --- Normalise comment safely ---
                special_df["comment_clean"] = (
                special_df["comment"]
                .fillna("")
                .str.lower()
                .str.strip()
                )
                # --- Classify manufacturer ---
                def classify_switch(comment):
                    if not isinstance(comment, str):
                        return "Unknown"
                    comment = comment.lower()
                    if re.search(r"\bsoule\b", comment):
                        return "Soule"
                    elif re.search(r"\bnoja\b", comment):
                        return "Noja"
                    else:
                        return "Unknown"

                special_df["Manufacturer"] = special_df["comment_clean"].apply(classify_switch)

                # --- Aggregate ---
                special_summary = (special_df.groupby(["item", "Manufacturer"], as_index=False)["Quantity_used"]
                                   .sum().rename(columns={"item": "Description","Quantity_used": "Total Quantity","Manufacturer": "Comment",}))

            else:
                special_summary = pd.DataFrame(columns=["Description", "Total Quantity", "Comment"])

            # Append special item summary (multiple rows per comment)
            final_summary = pd.concat([general_summary, special_summary], ignore_index=True, sort=False)

            # Write summary sheet
            final_summary.to_excel(writer, sheet_name="Summary", index=False, startrow=1)
            ws_summary = writer.book["Summary"]

        # ---- Formatting styles ----
        header_font = Font(bold=True, size=16)
        header_fill = PatternFill(start_color="00CCFF", end_color="00CCFF", fill_type="solid")
        thin_side = Side(style="thin")
        medium_side = Side(style="medium")
        thick_side = Side(style="thick")
        light_grey_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

        # AFTER ✅
        for sheet in [ws, ws_summary]:
            sheet.row_dimensions[1].height = 90   # logo row

        # ---- Load & resize images ----
        IMG_HEIGHT = 120
        IMG_WIDTH_SMALL = 120
        IMG_WIDTH_LARGE = IMG_WIDTH_SMALL * 3  # 🔹 3× wider

        img1 = XLImage("Images/GaeltecImage.png")
        img2 = XLImage("Images/SPEN.png")

        img1.width = IMG_WIDTH_SMALL
        img1.height = IMG_HEIGHT

        img2.width = IMG_WIDTH_LARGE
        img2.height = IMG_HEIGHT

        # Position images (row 1)
        img1.anchor = "B1"
        img2.anchor = "A1"

        ws.add_image(img1)
        ws.add_image(img2)

        # Same for Summary
        img1_s = XLImage("Images/GaeltecImage.png")
        img2_s = XLImage("Images/SPEN.png")

        img1_s.width = IMG_WIDTH_SMALL
        img1_s.height = IMG_HEIGHT
        img1_s.anchor = "A1"

        img2_s.width = IMG_WIDTH_LARGE
        img2_s.height = IMG_HEIGHT
        img2_s.anchor = "B1"

        ws_summary.add_image(img1_s)
        ws_summary.add_image(img2_s)


        # ---- Formatting (unchanged style) ----
        for sheet in [ws, ws_summary]:
            max_col = sheet.max_column
            max_row = sheet.max_row

            # HEADER → ROW 2 ✅
            for col_idx, cell in enumerate(sheet[2], start=1):
                cell.font = header_font
                cell.fill = header_fill
                sheet.column_dimensions[get_column_letter(col_idx)].width = 60 if col_idx == 1 else 20
                cell.border = Border(
                    left=thick_side if col_idx == 1 else medium_side,
                    right=thick_side if col_idx == max_col else medium_side,
                    top=thick_side,
                    bottom=thick_side
                )

            # DATA ROWS → START ROW 3 ✅
            for row_idx in range(3, max_row + 1):
                fill = light_grey_fill if row_idx % 2 == 1 else white_fill
                for col_idx in range(1, max_col + 1):
                    cell = sheet.cell(row=row_idx, column=col_idx)
                    cell.fill = fill
                    cell.border = Border(
                        left=thin_side,
                        right=thin_side,
                        top=thin_side,
                        bottom=thin_side
                    )

    # ---- Download button ----
    buffer_agg.seek(0)
    st.download_button(
        label="📥 Download Excel (Output Details)",
        data=buffer_agg,
        file_name="Gaeltec_Output.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

else:
    st.info("Project or Segment Code columns not found in the data.")

# -------------------------------
# Jobs per Team per Day
# -------------------------------
if {'datetouse_dt', 'team_name', 'total'}.issubset(filtered_df.columns):
    team_df = (
        filtered_df
        .dropna(subset=['datetouse_dt', 'team_name'])
        .groupby(['datetouse_dt', 'team_name'], as_index=False)['total']
        .sum()
    )

    fig_team = px.line(
        team_df,
        x='datetouse_dt',
        y='total',
        color='team_name',
        markers=True,
        title="Jobs per Team per Day"
    )
    st.plotly_chart(fig_team, use_container_width=True)


    # -------------------------------
    # Revenue per Project (Excel Export)
    # -------------------------------
    if not filtered_df.empty and 'project' in filtered_df.columns and 'total' in filtered_df.columns:
        revenue_per_project = (
            filtered_df
            .groupby('project', as_index=False)['total']
            .sum()
            .sort_values('total', ascending=False)
       )

        revenue_per_project.rename(
            columns={'total': 'Revenue (£)'},
            inplace=True
        )
    else:
        revenue_per_project = pd.DataFrame()
    
    if not filtered_df.empty and 'team_name' in filtered_df.columns and 'total' in filtered_df.columns:
        revenue_per_team = (
            filtered_df
            .groupby('team_name', as_index=False)['total']
            .sum()
            .sort_values('total', ascending=False)
        )

        revenue_per_team.rename(
            columns={'team_name': 'Team', 'total': 'Revenue (£)'},
            inplace=True
        )
    else:
        revenue_per_team = pd.DataFrame()

    if not revenue_per_project.empty or not revenue_per_team.empty:
        excel_file = to_excel(revenue_per_project, revenue_per_team)
        st.download_button(
            label="📥 Download Revenue Summary (Excel)",
            data=excel_file,
            file_name=f"revenue_summary_{date_range_str}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    else:
        st.info("No revenue data available for export.")
    
    # Display Project and completion
    col_top_left, col_top_right = st.columns([1, 1])
    # Project Completion
    with col_top_left:
        st.markdown("<h3 style='text-align:center; color:white;'>Projects Distribution</h3>", unsafe_allow_html=True)
        # --- Top-right Pie Chart: Projects Distribution ---
        try:
            if 'filtered_df' in locals() and not filtered_df.empty and 'project' in filtered_df.columns:
                
                # Count projects and get top projects
                project_counts = filtered_df['project'].value_counts().reset_index()
                project_counts.columns = ['Project', 'total']
                
                # If too many projects, group smaller ones into "Other"
                if len(project_counts) > 8:
                    top_projects = project_counts.head(7)
                    other_count = project_counts['total'].iloc[7:].sum()
                    other_row = pd.DataFrame({'Project': ['Other'], 'total': [other_count]})
                    project_data = pd.concat([top_projects, other_row], ignore_index=True)
                else:
                    project_data = project_counts
                
                # Create pie chart
                fig_projects = px.pie(
                    project_data,
                    names='Project',
                    values='total',
                    title="",
                    hole=0.4
                )
                fig_projects.update_traces(
                    textinfo='percent+label',
                    textfont_size=14,
                    marker=dict(line=dict(color='#000000', width=1))
                )
                fig_projects.update_layout(
                    title_text="",
                    title_font_size=16,
                    font=dict(color='white'),
                    paper_bgcolor='rgba(0,0,0,0)',
                    plot_bgcolor='rgba(0,0,0,0)',
                    showlegend=False,
                    annotations=[dict(text=f'Total<br>{len(filtered_df)}', x=0.5, y=0.5, font_size=16, showarrow=False)]
                )
                
                st.plotly_chart(fig_projects, use_container_width=True)
                
            else:
                st.info("No project data available for the selected filters.")
                
        except Exception as e:
            st.warning(f"Could not generate projects pie chart: {e}")

    # Works total
    with col_top_right:
        # Left side: Projects & Segments Overview and Works Complete pie chart
        col_left_top, col_left_bottom = st.columns([1, 1])
        
        with col_left_top:
            st.markdown("<h3 style='color:white;'>Projects & Segments Overview</h3>", unsafe_allow_html=True)

            if 'project' in filtered_df.columns and 'segmentcode' in filtered_df.columns:
                projects = filtered_df['project'].dropna().unique()
                if len(projects) == 0:
                    st.info("No projects found for the selected filters.")
                else:
                    for proj in sorted(projects):
                        segments = filtered_df[filtered_df['project'] == proj]['segmentcode'].dropna().unique()
                    
                        # Use expander to make segment list scrollable
                        with st.expander(f"Project: {proj} ({len(segments)} segments)"):
                            if len(segments) > 0:
                                # Scrollable container for segments
                                st.markdown(
                                    "<div style='max-height:150px; overflow-y:auto; padding:5px; border:1px solid #444;'>"
                                    + "<br>".join(segments.astype(str))
                                    + "</div>",
                                    unsafe_allow_html=True
                                )
                            else:
                                st.write("No segment codes for this project.")
            else:
                st.info("Project or Segment Code columns not found in the data.")

    # -----------------------------
    # Streamlit download button
    # -----------------------------

# ---- Streamlit download button ----
    if 'filtered_df' in locals() and not filtered_df.empty:
        excel_file = generate_excel_styled_multilevel(
            filtered_df,
            poles_df if 'poles_df' in locals() else None)
        st.download_button(
            label="📥 High level planning & Poles Excel",
            data=excel_file,
            file_name=f"High level planning_{date_range_str}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
    # -------------------------------
    # --- Map Section ---
    # -------------------------------
    col_map, col_desc = st.columns([2, 1])
    with col_map:
        st.header("🗺️ Regional Map View")
        folder_path = r"Maps"
        file_list = glob.glob(os.path.join(folder_path, "*.json"))

        if not file_list:
            st.error(f"No JSON files found in folder: {folder_path}")
        else:
            gdf_list = [gpd.read_file(file) for file in file_list]
            combined_gdf = gpd.GeoDataFrame(pd.concat(gdf_list, ignore_index=True), crs=gdf_list[0].crs)

            if "region" in filtered_df.columns:
                active_regions = filtered_df["region"].dropna().unique().tolist()
                wards_to_select = []
                for region in active_regions:
                    if region in mapping_region:
                        wards_to_select.extend(mapping_region[region])
                    else:
                        wards_to_select.append(region)
                wards_to_select = list(set(wards_to_select))
                areas_of_interest = combined_gdf[combined_gdf["WD13NM"].isin(wards_to_select)]
            else:
                areas_of_interest = pd.DataFrame()

            if not areas_of_interest.empty:
                areas_of_interest["geometry_simplified"] = areas_of_interest.geometry.simplify(tolerance=0.01)
                centroid = areas_of_interest.geometry_simplified.centroid.unary_union.centroid

                # Red flag
                flag_data = pd.DataFrame({"lon": [centroid.x], "lat": [centroid.y], "icon_name": ["red_flag"]})
                icon_mapping = {
                    "red_flag": {
                        "url": "https://upload.wikimedia.org/wikipedia/commons/thumb/3/3e/Red_flag_icon.svg/128px-Red_flag_icon.png",
                        "width": 128, "height": 128, "anchorY": 128
                    }
                }

                polygon_layer = pdk.Layer(
                    "GeoJsonLayer",
                    areas_of_interest["geometry_simplified"].__geo_interface__,
                    stroked=True,
                    filled=True,
                    get_fill_color=[160, 120, 80, 200],
                    get_line_color=[0, 0, 0],
                    pickable=True
                )

                flag_layer = pdk.Layer(
                    "IconLayer",
                    data=flag_data,
                    get_icon="icon_name",
                    get_size=4,
                    size_scale=15,
                    get_position='[lon, lat]',
                    pickable=True,
                    icon_mapping=icon_mapping
                )

                view_state = pdk.ViewState(latitude=centroid.y, longitude=centroid.x, zoom=8, pitch=0)

                st.pydeck_chart(
                    pdk.Deck(
                        layers=[polygon_layer, flag_layer],
                        initial_view_state=view_state,
                        map_style="mapbox://styles/mapbox/outdoors-v11"
                    )
                )
            else:
                st.info("No matching regions found for the selected filters.")


    with col_desc:
        st.markdown("<h3 style='color:white;'>Weather</h3>", unsafe_allow_html=True)
        
        # --- Scottish Weather Widget ---
        try:
            # Get API key from secrets
            api_key = st.secrets.get("d4d09fcf1373f72c30b970fb20d51fd9")
            
            if not api_key:
                st.info("Weather API key not configured")
            else:
                # Location selector
                location = st.selectbox(
                    "Select Location",
                    ["Ayrshire", "Lanarkshire", "Glasgow", "Edinburgh"],
                    index=0,
                    key="weather_location"
                )
                
                if st.button("Refresh Weather", key="refresh_weather"):
                    st.rerun()
                
                # Get current weather
                weather_data = get_scottish_weather(api_key, location)
                
                if weather_data:
                    # Display weather information
                    temp = weather_data['main']['temp']
                    feels_like = weather_data['main']['feels_like']
                    humidity = weather_data['main']['humidity']
                    wind_speed = weather_data['wind']['speed']
                    description = weather_data['weather'][0]['description'].title()
                    icon_code = weather_data['weather'][0]['icon']
                    
                    # Weather icon and description
                    col_icon, col_desc = st.columns([1, 2])
                    with col_icon:
                        st.image(f"http://openweathermap.org/img/wn/{icon_code}@2x.png", width=50)
                    with col_desc:
                        st.write(f"**{description}**")
                    
                    # Weather metrics
                    st.metric("Temperature", f"{temp}°C", f"Feels like {feels_like}°C")
                    st.metric("Humidity", f"{humidity}%")
                    st.metric("Wind Speed", f"{wind_speed} m/s")
                    
                    # Construction impact assessment
                    st.markdown("---")
                    st.markdown("**Construction Impact:**")
                    impact = assess_construction_impact(weather_data)
                    st.write(impact)
                else:
                    st.error("Failed to fetch weather data")
                    
        except Exception as e:
            st.warning(f"Could not load weather information: {e}")


# -------------------------------
# --- Mapping Bar Charts + Drill-down + Excel Export ---
# -------------------------------
    st.header("🪵 Materials")
    convert_to_miles = st.checkbox("Convert Equipment/Conductor Length to Miles")

    categories = [
        ("Poles 🪵", pole_keys, "Quantity"),
        ("Poles _erected 🪵", pole_erected_keys, "Quantity"),
        ("Poles _replaced 🪵", poles_replaced_keys, "Quantity"),
        ("Transformers ⚡🏭", transformer_keys, "Quantity"),
        ("Conductors", conductor_keys, "Length (Km)"),
        ("Conductors_2", conductor_2_keys, "Length (Km)"),
        ("Equipment", equipment_keys, "Quantity"),
        ("Insulators", insulator_keys, "Quantity"),
        ("LV Joints (Kits)", lv_joint_kit_keys, "Quantity"),
        ("LV Joint Modules", lv_joint_module_keys, "Quantity"),
        ("HV Joints / Terminations ⚡", hv_joint_termination_keys, "Quantity"),
        ("Cable Accessories 🔌", cable_accessory_keys, "Quantity"),
        ("Foundation & Steelwork 🏗️", foundation_steelwork_keys, "Quantity")
    ]

    def sanitize_sheet_name(name: str) -> str:
        name = str(name)
        name = re.sub(r'[:\\/*?\[\]\n\r]', '_', name)
        name = re.sub(r'[^\x00-\x7F]', '_', name)  # remove Unicode like m²
        return name[:31]


    for cat_name, keys, y_label in categories:

        # Only process if columns exist
        if 'item' not in filtered_df.columns or 'mapped' not in filtered_df.columns:
            st.warning("Missing required columns: item / mapped")
            continue
            
        # Build regex pattern for this category’s keys
        pattern = '|'.join([re.escape(k) for k in keys.keys()])

        mask = filtered_df['item'].astype(str).str.contains(pattern, case=False, na=False)
        sub_df = filtered_df[mask]

        if sub_df.empty:
            st.info(f"No data found for {cat_name}")
            continue

        # Aggregate
        if 'qsub' in sub_df.columns:
            sub_df['qsub_clean'] = pd.to_numeric(
                sub_df['qsub'].astype(str).str.replace(" ", "").str.replace(",", ".", regex=False),
                errors='coerce'
            )
            bar_data = sub_df.groupby('mapped')['qsub_clean'].sum().reset_index()
            bar_data.columns = ['Mapped', 'Total']
        else:
            bar_data = sub_df['mapped'].value_counts().reset_index()
            bar_data.columns = ['Mapped', 'Total']

        # Divide Conductors_2 by 1000
        if cat_name == "Conductors_2":
            bar_data['Total'] = bar_data['Total'] / 1000

        # Convert conductor units if needed
        y_axis_label = y_label
        if cat_name in ["Conductors", "Conductors_2"] and convert_to_miles:
            bar_data['Total'] = bar_data['Total'] * 0.621371
            y_axis_label = "Length (Miles)"

        # Compute grand total for the category
        grand_total = bar_data['Total'].sum()

        # Update Streamlit subheader with total
        st.subheader(f"🔹 {cat_name} — Total: {grand_total:,.2f}")

        # Draw the bar chart
        # FIX: Use go.Figure with explicit data types
        fig = go.Figure(data=[
            go.Bar(
                x=bar_data['Mapped'].astype(str).tolist(),
                y=bar_data['Total'].astype(float).tolist(),
                text=bar_data['Total'].astype(float).tolist(),
                texttemplate='%{y:,.1f}',
                textposition='outside'
            )
        ])

        fig.update_layout(
            title=f"{cat_name} Overview",
            xaxis_title="Mapping",
            yaxis_title=y_axis_label
        )
        
        # Add background colors separately
        fig.update_layout(
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            yaxis=dict(
                gridcolor='rgba(255,255,255,0.3)'  # Semi-transparent white grid
            )
        )

        # Display the chart
        st.plotly_chart(fig, use_container_width=True, height=500)

        # COLLAPSIBLE BUTTONS SECTION
        with st.expander("🔍 Click to explore more information", expanded=False):
            st.subheader("Select Mapping to Drill-down:")
            
            # Option 1: Buttons in columns
            cols = st.columns(3)  # 3 buttons per row
            
            for idx, mapping_value in enumerate(bar_data['Mapped']):
                col_idx = idx % 3  # Which column to use (0, 1, or 2)
                
                with cols[col_idx]:
                    button_key = f"btn_{cat_name}_{mapping_value}_{idx}"
                    
                    if st.button(f"📊 {mapping_value}", key=button_key, use_container_width=True):
                        st.session_state[f"selected_{cat_name}"] = mapping_value
                        st.rerun()  # Refresh to show the details immediately

        # Check if a mapping was selected
        selected_mapping = st.session_state.get(f"selected_{cat_name}")
        
        if selected_mapping:
            st.subheader(f"Details for: **{selected_mapping}**")
            
            # Add a button to clear the selection
            if st.button("❌ Clear Selection", key=f"clear_{cat_name}"):
                del st.session_state[f"selected_{cat_name}"]
                st.rerun()
            
            selected_rows = sub_df[sub_df['mapped'] == selected_mapping].copy()
            selected_rows.columns = selected_rows.columns.str.strip().str.lower()
            selected_rows = selected_rows.loc[:, ~selected_rows.columns.duplicated()]

            if 'datetouse' in selected_rows.columns:
                selected_rows['datetouse_display'] = pd.to_datetime(
                    selected_rows['datetouse'], errors='coerce'
                ).dt.strftime("%d/%m/%Y")
                selected_rows.loc[selected_rows['datetouse'].isna(), 'datetouse_display'] = "Unplanned"


            # Your original approach but working:
            extra_cols = ['poling team','team_name','shire','project','projectmanager','segmentcode','segmentdesc', 'material_code' ,'pid_ohl_nr', 'sourcefile' ]
            
            # Rename first
            selected_rows = selected_rows.rename(columns={
                "poling team": "code", 
                "team_name": "team lider"
            })

            # Update the extra_cols list to use new names
            extra_cols = [c if c != "poling team" else "code" for c in extra_cols]
            extra_cols = [c if c != "team_name" else "team lider" for c in extra_cols]


            # Filter to only existing columns
            extra_cols = [c for c in extra_cols if c in selected_rows.columns]
            # DEBUG: show the final columns being used
            st.write("🔹 Information Resumed:")
            # Create display date
            if 'datetouse' in selected_rows.columns:
                selected_rows['datetouse_display'] = pd.to_datetime(
                    selected_rows['datetouse'], errors='coerce'
                ).dt.strftime("%d/%m/%Y")
                selected_rows.loc[selected_rows['datetouse'].isna(), 'datetouse_display'] = "Unplanned"

            # 🔥 RENAME FOR DISPLAY
            selected_rows = selected_rows.rename(columns=column_rename_map)

            display_cols = ['Output','Quantity','material_code','pole','Date','District','project','Project Manager','Circuit','Segment','team lider','PID', 'sourcefile']
            display_cols = [c for c in display_cols if c in selected_rows.columns]
        

            if not selected_rows.empty:
                st.dataframe(selected_rows[display_cols], use_container_width=True)
                st.write(f"**Total records:** {len(selected_rows)}")
    
                if 'qsub_clean' in selected_rows.columns:
                    total_qsub = selected_rows['qsub_clean'].sum()
                    st.write(f"Total QSUB: {total_qsub:,.2f}")
            else:
                st.info("No records found for this selection")
                
            # Excel Export - Aggregated
            buffer_agg = BytesIO()
            with pd.ExcelWriter(buffer_agg, engine='openpyxl') as writer:
                aggregated_df = pd.DataFrame()
                for bar_value in bar_data['Mapped']:
                    df_bar = sub_df[sub_df['mapped'] == bar_value].copy()
                    df_bar = df_bar.loc[:, ~df_bar.columns.duplicated()]
                    if 'datetouse' in df_bar.columns:
                        df_bar['datetouse_display'] = pd.to_datetime(df_bar['datetouse'], errors='coerce')
                        df_bar['datetouse_display'] = df_bar['datetouse'].dt.strftime("%d/%m/%Y")
                        df_bar.loc[df_bar['datetouse'].isna(), 'datetouse_display'] = "Unplanned"

                    # 🔥 Rename columns BEFORE selecting
                    df_bar = df_bar.rename(columns=column_rename_map)

                    cols_to_include = ['Output','Quantity','material_code','pole','Date','District','project','Project Manager','Circuit','Segment','team lider','PID', 'sourcefile']
                    cols_to_include = [c for c in cols_to_include if c in df_bar.columns]
                    df_bar = df_bar[cols_to_include]

                    aggregated_df = pd.concat([aggregated_df, df_bar], ignore_index=True)

                aggregated_df.to_excel(writer, sheet_name='Aggregated', index=False)
                # Access the worksheet
                ws = writer.book['Aggregated']
                ws.insert_rows(1)
                # ---- Header style ----
                # ---- Formatting styles ----
                header_font = Font(bold=True, size=16)
                header_fill = PatternFill(start_color="00CCFF", end_color="00CCFF", fill_type="solid")
                thin_side = Side(style="thin")
                medium_side = Side(style="medium")
                thick_side = Side(style="thick")
                light_grey_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
                white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

                # AFTER ✅
                for sheet in [ws]:
                    sheet.row_dimensions[1].height = 90   # logo row

                # ---- Load & resize images ----
                IMG_HEIGHT = 120
                IMG_WIDTH_SMALL = 120
                IMG_WIDTH_LARGE = IMG_WIDTH_SMALL * 3  # 🔹 3× wider

                img1 = XLImage("Images/GaeltecImage.png")
                img2 = XLImage("Images/SPEN.png")

                img1.width = IMG_WIDTH_SMALL
                img1.height = IMG_HEIGHT

                img2.width = IMG_WIDTH_LARGE
                img2.height = IMG_HEIGHT

                # Position images (row 1)
                img1.anchor = "B1"
                img2.anchor = "A1"

                ws.add_image(img1)
                ws.add_image(img2)


                # ---- Formatting (unchanged style) ----
                for sheet in [ws]:
                    max_col = sheet.max_column
                    max_row = sheet.max_row

                    # HEADER → ROW 2 ✅
                    for col_idx, cell in enumerate(sheet[2], start=1):
                        cell.font = header_font
                        cell.fill = header_fill
                        sheet.column_dimensions[get_column_letter(col_idx)].width = 60 if col_idx == 1 else 20
                        cell.border = Border(
                            left=thick_side if col_idx == 1 else medium_side,
                            right=thick_side if col_idx == max_col else medium_side,
                            top=thick_side,
                            bottom=thick_side
                        )

                    # DATA ROWS → START ROW 3 ✅
                    for row_idx in range(3, max_row + 1):
                        fill = light_grey_fill if row_idx % 2 == 1 else white_fill
                        for col_idx in range(1, max_col + 1):
                            cell = sheet.cell(row=row_idx, column=col_idx)
                            cell.fill = fill
                            cell.border = Border(
                                left=thin_side,
                                right=thin_side,
                                top=thin_side,
                                bottom=thin_side
                            )

            buffer_agg.seek(0)
            st.download_button(
                f"📥 Download Excel (Aggregated): {cat_name} Details",
                buffer_agg,
                file_name=f"{cat_name}_Details_Aggregated.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

            # Excel Export - Separate Sheets
            buffer_sep = BytesIO()
            with pd.ExcelWriter(buffer_sep, engine='openpyxl') as writer:
                for bar_value in bar_data['Mapped']:
                    df_bar = sub_df[sub_df['mapped'] == bar_value].copy()
                    df_bar = df_bar.loc[:, ~df_bar.columns.duplicated()]
                    if 'datetouse' in df_bar.columns:
                        df_bar['datetouse_display'] = pd.to_datetime(
                            df_bar['datetouse'], errors='coerce'
                        )
                        df_bar.loc[df_bar['datetouse'].isna(), 'datetouse_display'] = "Unplanned"

                    cols_to_include = ['mapped', 'datetouse_display','qsub'] + extra_cols
                    cols_to_include = [c for c in cols_to_include if c in df_bar.columns]
                    df_bar = df_bar[cols_to_include]

                    sheet_name = sanitize_sheet_name(bar_value)
                    df_bar.to_excel(writer, sheet_name=sheet_name, index=False)

            buffer_sep.seek(0)
            st.download_button(
                f"📥 Download Excel (Separated): {cat_name} Details",
                buffer_sep,
                file_name=f"{cat_name}_Details_Separated.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

st.header("🛠️ Works")

if misc_df is not None and not misc_df.empty and not filtered_df.empty:
    # -----------------------------
    # Data preparation
    # -----------------------------
    filtered_df['item'] = filtered_df['item'].astype(str)
    misc_df['column_1'] = misc_df['column_1'].astype(str)
    misc_df['column_2'] = misc_df['column_2'].astype(str)

    # Keep only poles with valid values
    poles_df = filtered_df[
        filtered_df['pole'].notna() & 
        (filtered_df['pole'].astype(str).str.lower() != "nan")
    ].copy()

    # Merge filtered_df with misc_df to get Work instructions
    poles_df = poles_df.merge(
        misc_df[['column_1', 'column_2']],
        how='left',
        left_on='item',
        right_on='column_1'
    )

    poles_df.rename(columns={'column_2': 'Work instructions'}, inplace=True)

    # Keep only rows with valid instructions, comments, and team_name
    poles_df_clean = poles_df.dropna(subset=['Work instructions', 'comment', 'team_name'])[
        ['pole', 'segmentcode', 'Work instructions', 'comment', 'team_name']
    ]

    # -----------------------------
    # 🔘 Segment selector
    # -----------------------------
    segment_options = ['All'] + sorted(poles_df_clean['segmentcode'].dropna().astype(str).unique())
    selected_segment = st.selectbox("Select a segment code:", segment_options)

    if selected_segment != 'All':
        poles_df_view = poles_df_clean[poles_df_clean['segmentcode'].astype(str) == selected_segment]
    else:
        poles_df_view = poles_df_clean.copy()

    # -----------------------------
    # 🎯 Pole selector (Cascading)
    # -----------------------------
    pole_options = sorted(poles_df_view['pole'].dropna().astype(str).unique())
    selected_pole = st.selectbox("Select a pole to view details:", ["All"] + pole_options)

    # Filter by selected pole
    if selected_pole != "All":
        poles_df_view = poles_df_view[poles_df_view['pole'].astype(str) == selected_pole]

    # Display pole details if one is selected
    if selected_pole != "All" and not poles_df_view.empty:
        st.write(f"Details for pole **{selected_pole}**:")
        st.dataframe(poles_df_view)

    # -----------------------------
    # 📊 Pie chart (Works breakdown)
    # -----------------------------
    if not poles_df_view.empty:
        # Count work instructions and remove NaN / empty strings
        work_data = (
            poles_df_view['Work instructions']
            .astype(str)
            .replace('nan', pd.NA)
            .dropna()
            .value_counts()
            .reset_index()
        )
        work_data.columns = ['Work instructions', 'total']

        if not work_data.empty:
            fig_work = px.pie(
                work_data,
                names='Work instructions',
                values='total',
                hole=0.4
            )
            fig_work.update_traces(textinfo='percent+label', textfont_size=16)
            fig_work.update_layout(
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                showlegend=False
            )
            st.plotly_chart(fig_work, use_container_width=True)
        else:
            st.info("No valid work instructions available for the selected filters.")
    else:
        st.info("No data available for the selected filters.")

    # -----------------------------
    # 📄 Word export
    # -----------------------------
    if not poles_df_view.empty:
        word_file = poles_to_word(poles_df_view)
        st.download_button(
            label="⬇️ Download Work Instructions (.docx)",
            data=word_file,
            file_name="Pole_Work_Instructions.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

else:
    st.info("No work data available.")

# -----------------------------
# General summary placeholder
# -----------------------------
general_summary = pd.DataFrame(
    columns=["Description", "Total Quantity", "Comment"]
)
